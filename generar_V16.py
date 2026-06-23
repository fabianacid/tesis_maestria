"""
Genera TI_Cid_Fabiana_V16.docx a partir de V15:
  1. Corrige la portada (institución, director, ciudad)
  2. Inserta un índice (TOC) entre Dedicatoria y Capítulo 1
  3. Copia las imágenes de "Pantallas 22.06.26.docx" a la sección 4.3.5
"""

import io
import shutil
import lxml.etree as etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches

SRC      = "TI_Cid_Fabiana_V15.docx"
PANTALLAS = "Pantallas 22.06.26.docx"
DST      = "TI_Cid_Fabiana_V16.docx"

# ── 1. Copiar V15 → V16 ────────────────────────────────────────────────────
shutil.copy2(SRC, DST)
doc = Document(DST)


# ── 2. Corregir portada ────────────────────────────────────────────────────
# Párrafo 4: institución
p4 = doc.paragraphs[4]
for run in p4.runs:
    run.text = ""
p4.runs[0].text = "Tesis Final de la Maestría en Finanzas\nUniversidad Nacional de Rosario (UNR)"

# Párrafo 6: director + ciudad (elimina co-director y jurados)
p6 = doc.paragraphs[6]
for run in p6.runs:
    run.text = ""
p6.runs[0].text = "Director: Ph.D. Luciano Machain (UNR)\n\nCiudad de Rosario, junio de 2026"

print("Portada corregida.")


# ── 3. Helpers para el TOC ─────────────────────────────────────────────────
def make_toc_field():
    """Párrafo con campo TOC de Word (se actualiza con F9)."""
    p = OxmlElement("w:p")

    r1 = OxmlElement("w:r")
    r1.append(OxmlElement("w:rPr"))
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    r1.append(fldBegin)
    p.append(r1)

    r2 = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instrText)
    p.append(r2)

    r3 = OxmlElement("w:r")
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    r3.append(fldSep)
    p.append(r3)

    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "(Actualice el campo: Ctrl+A → F9 o clic derecho → Actualizar campo)"
    r4.append(t)
    p.append(r4)

    r5 = OxmlElement("w:r")
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    r5.append(fldEnd)
    p.append(r5)

    return p


def make_heading(text, level=1):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), f"Heading{level}")
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_page_break():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


# ── 4. Insertar TOC entre Dedicatoria y Capítulo 1 ────────────────────────
body = doc.element.body
all_elems = list(body)

# Mapear índice de párrafo → índice en body
para_to_body = {}
pi = 0
for i, elem in enumerate(all_elems):
    if elem.tag == qn("w:p"):
        para_to_body[pi] = i
        pi += 1

# Encontrar primer Heading 1 con "Cap"
cap1_para = next(
    i for i, p in enumerate(doc.paragraphs)
    if p.style.name == "Heading 1" and "Cap" in p.text
)
cap1_body_idx = para_to_body[cap1_para]
ref = all_elems[cap1_body_idx]

body.insert(cap1_body_idx,     make_page_break())
body.insert(cap1_body_idx + 1, make_heading("Índice general", level=1))
body.insert(cap1_body_idx + 2, make_toc_field())
body.insert(cap1_body_idx + 3, make_page_break())

print(f"TOC insertado antes del párrafo {cap1_para} ('{doc.paragraphs[cap1_para].text[:40]}').")


# ── 5. Extraer imágenes de Pantallas ──────────────────────────────────────
def get_ordered_images(source_doc):
    images = []
    for para in source_doc.paragraphs:
        for elem in para._element.iter():
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "blip":
                rId = elem.get(qn("r:embed"))
                if rId and rId in source_doc.part.rels:
                    rel = source_doc.part.rels[rId]
                    if "image" in rel.reltype:
                        images.append((rel.target_part.blob, rel.target_part.content_type))
    return images

pantallas_doc = Document(PANTALLAS)
images = get_ordered_images(pantallas_doc)
print(f"Imágenes encontradas en Pantallas: {len(images)}")


# ── 6. Insertar imágenes en sección 4.3.5 ─────────────────────────────────
# Guardar y reabrir para que los índices de párrafo sean correctos
doc.save(DST)
doc = Document(DST)

# Encontrar 4.3.5 y el heading siguiente
section_435 = next(i for i, p in enumerate(doc.paragraphs) if "4.3.5" in p.text and "Capturas" in p.text)
insert_before = section_435 + 1
while insert_before < len(doc.paragraphs):
    if doc.paragraphs[insert_before].style.name.startswith("Heading"):
        break
    insert_before += 1

print(f"Insertando imágenes antes del párrafo {insert_before}: '{doc.paragraphs[insert_before].text[:50]}'")

# Reconstruir mapa body
body = doc.element.body
all_elems2 = list(body)
para_to_body2 = {}
pi2 = 0
for i, elem in enumerate(all_elems2):
    if elem.tag == qn("w:p"):
        para_to_body2[pi2] = i
        pi2 += 1

ref_elem = all_elems2[para_to_body2[insert_before]]

for idx, (blob, ct) in enumerate(images):
    # Crear párrafo de imagen al final del doc
    stream = io.BytesIO(blob)
    para = doc.add_paragraph()
    para.alignment = 1
    run = para.add_run()
    run.add_picture(stream, width=Inches(5.5))
    p_img_elem = para._element

    # Asignar IDs únicos
    for elem in p_img_elem.iter():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "docPr":
            elem.set("id", str(200 + idx))
            elem.set("name", f"Pantalla_{idx+1:02d}")
        elif tag == "cNvPr":
            elem.set("id", str(200 + idx))

    # Mover al lugar correcto
    body.remove(p_img_elem)
    ref_elem.addprevious(p_img_elem)

    # Pie de imagen
    cap_p = OxmlElement("w:p")
    cap_pPr = OxmlElement("w:pPr")
    cap_jc = OxmlElement("w:jc")
    cap_jc.set(qn("w:val"), "center")
    cap_pPr.append(cap_jc)
    cap_pStyle = OxmlElement("w:pStyle")
    cap_pStyle.set(qn("w:val"), "Caption")
    cap_pPr.append(cap_pStyle)
    cap_p.append(cap_pPr)
    cap_r = OxmlElement("w:r")
    cap_t = OxmlElement("w:t")
    cap_t.text = f"Figura: Captura de pantalla {idx+1} - interfaz del sistema (22/06/2026)"
    cap_r.append(cap_t)
    cap_p.append(cap_r)
    ref_elem.addprevious(cap_p)

    print(f"  Imagen {idx+1}/{len(images)} insertada.")


# ── 7. Guardar V16 ────────────────────────────────────────────────────────
doc.save(DST)
print(f"\nGuardado: {DST}")
print("Abrir en Word y presionar Ctrl+A -> F9 para actualizar el indice.")
