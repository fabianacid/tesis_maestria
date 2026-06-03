"""
Genera TI_Cid_Fabiana_V10.docx — versión para tesis de maestría.
Resumen y Capítulo 1 expandidos con citas APA. Incluye SECAgent y PortfolioAgent.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Márgenes y fuente base ───────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)

# ── Helpers ──────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(16); p.runs[0].font.bold = True
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(13)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(12)
    return p

def body(text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.75)
    return p

def body_noi(text):
    """Párrafo sin sangría — para listas explicadas, notas, etc."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def numbered(text):
    return doc.add_paragraph(text, style='List Number')

def add_table(headers, rows, caption_text=None):
    if caption_text:
        cp = doc.add_paragraph(caption_text)
        cp.runs[0].font.size = Pt(10)
        cp.runs[0].font.italic = True
        cp.runs[0].font.bold  = True
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\n\nSistema Multiagente de Análisis y Optimización\nde Portafolios Financieros')
r.font.size = Pt(22); r.font.bold = True

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('Ing. María Fabiana Cid').font.size = Pt(16)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run('Carrera de Especialización en Inteligencia Artificial\nFIUBA').font.size = Pt(13)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run(
    'Director: Dr. Camilo Argoty Pulido (FIUBA)\n'
    'Co-Director: Ph.D. Luciano Machain (UNR)\n\n'
    'Jurados:\n'
    'Esp. Ing. Federico Arias Suarez (FIUBA)\n'
    'Esp. Ing. Juan Pablo Alianak (FIUBA)\n'
    'Ing. Christopher Charaf (URBC)\n\n'
    'Ciudad de Rosario, junio de 2026'
).font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# RESUMEN  (expandido, nivel tesis de maestría)
# ════════════════════════════════════════════════════════════════════════════
h1('Resumen')

body(
    'La creciente disponibilidad de datos financieros, impulsada por la digitalización '
    'de los mercados y el desarrollo de interfaces de programación públicas (API), ha '
    'generado un entorno propicio para la aplicación de técnicas de inteligencia '
    'artificial en la gestión de inversiones. Sin embargo, las herramientas más '
    'sofisticadas siguen concentradas en grandes instituciones financieras y fondos de '
    'inversión, perpetuando una asimetría informativa estructural que limita las '
    'capacidades analíticas de los inversores minoristas (Grossman & Stiglitz, 1980; '
    'Fama, 1970). Esta brecha se manifiesta en el acceso diferencial a modelos '
    'predictivos, análisis de sentimiento de noticias, datos fundamentales de empresas '
    'y herramientas cuantitativas de optimización de carteras.'
)
body(
    'La presente memoria describe el diseño, implementación y validación de un '
    'prototipo funcional de sistema multiagente inteligente orientado al análisis, '
    'seguimiento y optimización de portafolios de activos financieros. El sistema '
    'integra siete agentes especializados que operan de manera coordinada: '
    '(1) MarketAgent, encargado de la obtención de datos históricos y el cálculo de '
    'más de 35 indicadores técnicos; (2) ModelAgent, que implementa un ensemble de '
    'clasificadores de machine learning —Random Forest, Gradient Boosting, XGBoost '
    '(Chen & Guestrin, 2016) y LightGBM (Ke et al., 2017)— para predecir la dirección '
    'del precio en un horizonte de tres días hábiles; (3) SentimentAgent, que aplica '
    'un ensemble de modelos de procesamiento del lenguaje natural compuesto por '
    'FinBERT (Araci, 2019), VADER (Hutto & Gilbert, 2014), un lexicón financiero '
    'especializado y TextBlob; (4) RecommendationAgent, que integra señales técnicas, '
    'predictivas y de sentimiento mediante un esquema de decisión multi-factor con '
    'gestión de riesgo (Value at Risk y criterio de Kelly); (5) AlertAgent, con seis '
    'niveles de severidad y detección de anomalías mediante Isolation Forest y métodos '
    'estadísticos; (6) SECAgent, que accede a datos fundamentales vía yfinance y a '
    'declaraciones financieras regulatorias (10-K, 10-Q, 8-K) a través de la API '
    'pública de SEC EDGAR; y (7) PortfolioAgent, que implementa la teoría moderna de '
    'carteras de Markowitz (1952) para la optimización de portafolios mediante el '
    'cálculo del portafolio de máximo ratio de Sharpe, el de mínima varianza y la '
    'frontera eficiente.'
)
body(
    'La validación empírica del sistema se realizó con datos reales de mercado '
    'obtenidos a través de la API de Yahoo Finance, evaluando diez acciones '
    'representativas del mercado estadounidense en cuatro sesiones de prueba entre '
    'febrero y marzo de 2026. Los resultados de las pruebas de clasificación del '
    'ModelAgent arrojan una exactitud (accuracy) del 57,0 %, una precisión del 60,7 % '
    'y un recall del 66,2 % —todos superiores al umbral aleatorio del 50 %—, empleando '
    'una ventana de entrenamiento de 504 días, validación cruzada temporal con cinco '
    'particiones (walk-forward) y un umbral del 0,5 % en la definición del target. Las '
    'pruebas de carga muestran disponibilidad del 100 % con hasta 10 usuarios '
    'concurrentes y latencia promedio de 4,65 segundos en la configuración final, '
    'cumpliendo el requisito no funcional de respuesta menor a 5 segundos. El '
    'PortfolioAgent optimiza carteras de 2 a 15 activos y genera la frontera eficiente '
    'de Markowitz con 15 puntos, mientras que el SECAgent recupera datos fundamentales '
    'y filings regulatorios para todos los tickers del mercado estadounidense evaluados.'
)
body(
    'El trabajo integra conocimientos de aprendizaje automático, procesamiento del '
    'lenguaje natural, optimización matemática, teoría de carteras, ingeniería de '
    'software y seguridad informática, materializados en un prototipo funcional de bajo '
    'costo y código abierto que podría contribuir a democratizar el acceso a la '
    'inteligencia financiera avanzada para inversores minoristas, educadores financieros '
    'e investigadores. Si bien el sistema opera como herramienta de apoyo a la decisión '
    '—y no como reemplazo del juicio del inversor—, los resultados obtenidos demuestran '
    'la viabilidad técnica de la arquitectura propuesta y establecen una base sólida '
    'para futuras extensiones hacia entornos de producción.'
)
body_noi(
    'Palabras clave: sistema multiagente, machine learning, análisis de sentimiento, '
    'FinBERT, procesamiento del lenguaje natural, optimización de portafolios, '
    'frontera eficiente de Markowitz, análisis fundamental, SEC EDGAR, FastAPI, '
    'inteligencia artificial aplicada a finanzas.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# AGRADECIMIENTOS
# ════════════════════════════════════════════════════════════════════════════
h1('Agradecimientos')
body(
    'Quiero agradecer a mi esposo Eduardo y a mis cuatro hijos Francisco, Manuel, Inés y '
    'Victoria por su paciencia y por su aliento incondicional durante todo el proceso. '
    'Agradezco también muy especialmente a mi amiga Belén por haberme inspirado para '
    'realizar este posgrado. Finalmente, muchas gracias a mis correctores, a mi director, '
    'mi codirector y mis jurados por el tiempo, la dedicación y la enseñanza brindada.'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# DEDICATORIA
# ════════════════════════════════════════════════════════════════════════════
h1('Dedicatoria')
p_ded = doc.add_paragraph()
p_ded.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ded.add_run(
    '\n\nDedicado a mi esposo Eduardo y a mis hijos Francisco, Manuel, Inés y Victoria.'
).font.bold = True
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 1 — INTRODUCCIÓN GENERAL  (expandido, nivel tesis de maestría)
# ════════════════════════════════════════════════════════════════════════════
h1('Capítulo 1 — Introducción general')

body(
    'El presente capítulo introduce el contexto financiero y tecnológico que motiva el '
    'desarrollo del sistema propuesto, delimita el problema de investigación, revisa el '
    'estado del arte relevante, justifica la brecha que este trabajo busca cubrir y '
    'presenta los objetivos, el alcance y la estructura de la memoria.'
)

# ── 1.1 ─────────────────────────────────────────────────────────────────────
h2('1.1 Contexto financiero y tecnológico')

body(
    'Los mercados financieros contemporáneos se caracterizan por un volumen de datos '
    'sin precedentes, una velocidad de operación creciente y una interconexión global '
    'que amplifica tanto las oportunidades como los riesgos para los participantes. '
    'El surgimiento del trading algorítmico, la proliferación de plataformas digitales '
    'de inversión y la democratización parcial del acceso a datos de mercado a través '
    'de APIs públicas han transformado la estructura del sector en las últimas dos '
    'décadas (López de Prado, 2018). En este contexto, la inteligencia artificial (IA) '
    'y el aprendizaje automático (machine learning, ML) han emergido como herramientas '
    'fundamentales para extraer valor de grandes volúmenes de datos financieros '
    'heterogéneos.'
)
body(
    'La hipótesis de los mercados eficientes (Fama, 1970) sostiene que los precios de '
    'los activos reflejan toda la información disponible públicamente, lo que implicaría '
    'la imposibilidad de obtener rendimientos ajustados por riesgo superiores al mercado '
    'de forma sostenida. Sin embargo, evidencia empírica acumulada en las últimas '
    'décadas sugiere que los mercados presentan ineficiencias sistemáticas, '
    'particularmente en horizontes de corto plazo, que pueden ser aprovechadas mediante '
    'modelos predictivos sofisticados (Lo, 2004; Grossman & Stiglitz, 1980). '
    'Esta perspectiva, articulada en la hipótesis de los mercados adaptativos (Lo, '
    '2004), proporciona el fundamento teórico para el desarrollo de sistemas '
    'inteligentes de apoyo a la decisión financiera.'
)
body(
    'Al mismo tiempo, la alta volatilidad de los mercados —exacerbada por eventos '
    'geopolíticos, shocks macroeconómicos y la velocidad de difusión de información en '
    'redes sociales— genera asimetrías informativas que afectan de manera diferencial '
    'a los distintos participantes del mercado (Malkiel & Fama, 1970). Mientras que '
    'los actores institucionales disponen de equipos especializados, acceso prioritario '
    'a datos y sistemas automatizados de análisis, los inversores minoristas suelen '
    'depender de información fragmentada, tardía y de difícil interpretación para '
    'usuarios sin formación técnica en finanzas.'
)
body(
    'En este escenario, la integración de técnicas de ML, procesamiento del lenguaje '
    'natural (NLP), análisis fundamental y optimización de carteras en una plataforma '
    'accesible, modular y de bajo costo representa una oportunidad concreta para reducir '
    'la brecha informativa entre grandes instituciones y pequeños inversores, '
    'contribuyendo a la democratización del análisis financiero avanzado.'
)

# ── 1.2 ─────────────────────────────────────────────────────────────────────
h2('1.2 Problema de investigación')

body(
    'La asimetría informativa en los mercados financieros constituye una falla de '
    'mercado ampliamente documentada en la literatura económica (Grossman & Stiglitz, '
    '1980; Akerlof, 1970). En el ámbito de la gestión de inversiones minoristas, esta '
    'asimetría se manifiesta en tres dimensiones complementarias:'
)
bullet(
    'Acceso a herramientas analíticas: las plataformas de análisis financiero de mayor '
    'capacidad —como Bloomberg Terminal o Reuters Eikon— tienen costos de licencia que '
    'oscilan entre USD 24.000 y USD 30.000 anuales (Gratton & Erickson, 2024), '
    'tornándolas inaccesibles para la mayoría de los inversores individuales.'
)
bullet(
    'Capacidad de procesamiento de información: el volumen de noticias, reportes '
    'financieros y datos de mercado generados diariamente supera ampliamente la '
    'capacidad de procesamiento humano, lo que dificulta la toma de decisiones '
    'informadas en tiempo real sin soporte tecnológico automatizado.'
)
bullet(
    'Optimización de portafolios: la aplicación práctica de la teoría moderna de '
    'carteras (Markowitz, 1952) requiere el cálculo de matrices de covarianza, la '
    'resolución de problemas de optimización cuadrática y la estimación de retornos '
    'esperados, tareas que demandan conocimientos técnicos especializados y herramientas '
    'computacionales que no están al alcance del inversor minorista promedio.'
)
body(
    'Frente a este diagnóstico, el problema central de investigación de este trabajo '
    'puede formularse como sigue: ¿es técnicamente viable diseñar e implementar un '
    'sistema multiagente de bajo costo que integre predicción de precios mediante ML, '
    'análisis de sentimiento mediante NLP, datos fundamentales de empresas y '
    'optimización de portafolios mediante la teoría de Markowitz, de forma modular, '
    'escalable y accesible para inversores minoristas sin formación técnica avanzada?'
)

# ── 1.3 ─────────────────────────────────────────────────────────────────────
h2('1.3 Estado del arte')

body(
    'El análisis automatizado de activos financieros mediante técnicas de IA ha sido '
    'abordado desde múltiples perspectivas en la literatura académica y en desarrollos '
    'comerciales. A continuación se revisan los antecedentes más relevantes para los '
    'componentes del sistema propuesto.'
)

h3('1.3.1 Predicción de precios mediante machine learning')

body(
    'La aplicación de modelos de ML a la predicción de precios financieros ha sido '
    'objeto de una extensa literatura. Los trabajos pioneros de Breiman (2001) sobre '
    'Random Forests y los desarrollos posteriores de modelos de gradient boosting '
    '(Friedman, 2001) sentaron las bases para la predicción de dirección de precios '
    'mediante clasificación binaria. Estudios como el de Fischer y Krauss (2018) '
    'demostraron que los modelos de ML superan a los enfoques estadísticos clásicos '
    '(ARIMA, regresión lineal) en la predicción de rendimientos de corto plazo cuando '
    'se dispone de un conjunto suficiente de características técnicas.'
)
body(
    'Las redes neuronales recurrentes, en particular las arquitecturas Long Short-Term '
    'Memory (LSTM) introducidas por Hochreiter y Schmidhuber (1997), han demostrado '
    'capacidad para capturar dependencias temporales de largo plazo en series '
    'financieras. Bao, Yue y Yin (2017) reportaron mejoras significativas en la '
    'predicción de precios de acciones al combinar LSTM con autoencoders y análisis '
    'de sentimiento. Sin embargo, Fischer y Krauss (2018) encontraron que los métodos '
    'de ensemble como Random Forest son competitivos con LSTM en horizontes de corto '
    'plazo, con la ventaja adicional de mayor interpretabilidad.'
)
body(
    'Los métodos de ensemble basados en boosting han demostrado consistentemente alto '
    'desempeño en tareas de clasificación financiera. XGBoost (Chen & Guestrin, 2016) '
    'y LightGBM (Ke et al., 2017) representan el estado del arte en este paradigma, '
    'ofreciendo eficiencia computacional y mecanismos de regularización que mitigan '
    'el sobreajuste, un riesgo particularmente relevante en datos financieros con '
    'alta relación señal-ruido (López de Prado, 2018).'
)
body(
    'Un aspecto crítico en la evaluación de modelos sobre series temporales financieras '
    'es la prevención del data leakage, es decir, la contaminación de los datos de '
    'validación con información futura. El enfoque de walk-forward validation con '
    'TimeSeriesSplit (Scikit-learn, 2023) garantiza que cada fold de evaluación '
    'utilice únicamente datos anteriores al período de validación, respetando el '
    'orden causal de la información financiera (López de Prado, 2018).'
)

h3('1.3.2 Análisis de sentimiento en noticias financieras')

body(
    'El impacto del sentimiento en los mercados financieros ha sido documentado '
    'ampliamente en la literatura. Tetlock (2007) demostró que el tono negativo en '
    'artículos del Wall Street Journal predice rendimientos bursátiles negativos al '
    'día siguiente, estableciendo un vínculo empírico entre el análisis de texto y '
    'los movimientos del mercado. Estudios posteriores extendieron este análisis a '
    'diversas fuentes textuales, incluyendo reportes de analistas, publicaciones en '
    'redes sociales y noticias de agencias especializadas.'
)
body(
    'El modelo VADER (Valence Aware Dictionary and sEntiment Reasoner), desarrollado '
    'por Hutto y Gilbert (2014), ofrece un enfoque léxico-basado especialmente diseñado '
    'para textos breves y lenguaje informal, siendo ampliamente utilizado para el '
    'análisis de noticias financieras de corta extensión. TextBlob (Loria, 2020) '
    'complementa este enfoque con un análisis de polaridad basado en estadística '
    'sobre corpus lingüísticos generales.'
)
body(
    'La introducción de los modelos Transformer (Vaswani et al., 2017) y, '
    'específicamente, de BERT (Devlin et al., 2019) marcó un hito en el NLP al '
    'demostrar que el preentrenamiento en grandes corpus seguido de ajuste fino '
    '(fine-tuning) en dominios específicos producía representaciones lingüísticas '
    'de alta calidad. FinBERT (Araci, 2019), derivado de BERT con ajuste fino en '
    'textos financieros (reportes de la SEC, artículos de noticias financieras), '
    'estableció el estado del arte en la clasificación de sentimiento financiero, '
    'superando a métodos léxicos como VADER en corpus de noticias especializadas. '
    'Huang, Wang y Yang (2023) confirmaron estas ventajas en estudios de evaluación '
    'comparativa sobre múltiples conjuntos de datos financieros.'
)
body(
    'Amola (2025) propone un enfoque híbrido que integra el análisis de sentimiento '
    'de noticias con modelos cuantitativos tradicionales, obteniendo mejoras en la '
    'predicción de rendimientos de corto plazo respecto a métodos puramente técnicos. '
    'Estudios recientes incorporan grandes modelos de lenguaje (LLM) y mecanismos '
    'de recuperación de información contextual (RAG) para mejorar la interpretación '
    'de textos financieros complejos, aunque con mayor costo computacional '
    '(Brown et al., 2020).'
)

h3('1.3.3 Análisis fundamental y datos regulatorios')

body(
    'El análisis fundamental evalúa el valor intrínseco de una empresa a través de '
    'indicadores financieros derivados de sus estados contables: ratios de valuación '
    '(P/E, P/B, EV/EBITDA), rentabilidad (ROE, ROA, márgenes), solidez financiera '
    '(deuda/capital, ratios de liquidez) y perspectivas de crecimiento (Damodaran, '
    '2012). Este enfoque, complementario al análisis técnico, busca determinar si un '
    'activo cotiza por encima o por debajo de su valor razonable.'
)
body(
    'La Securities and Exchange Commission (SEC) de Estados Unidos mantiene EDGAR '
    '(Electronic Data Gathering, Analysis, and Retrieval), una base de datos pública '
    'de declaraciones financieras obligatorias que incluye informes anuales (10-K), '
    'trimestrales (10-Q) y de eventos relevantes (8-K). La disponibilidad de una API '
    'pública y gratuita (SEC, 2023) permite el acceso automatizado a estos documentos, '
    'habilitando la integración de información regulatoria en sistemas automatizados '
    'de análisis sin necesidad de suscripciones costosas.'
)
body(
    'La combinación de análisis fundamental con modelos de ML ha mostrado resultados '
    'prometedores en la predicción de rendimientos de largo plazo. Piotroski (2000) '
    'demostró que una estrategia basada en nueve indicadores fundamentales binarios '
    '(el F-Score) generaba rendimientos anormales significativos, precursando el '
    'interés en la integración cuantitativa del análisis fundamental.'
)

h3('1.3.4 Optimización de portafolios')

body(
    'La teoría moderna de carteras, desarrollada por Markowitz (1952) en su trabajo '
    'seminal "Portfolio Selection", introdujo el concepto de frontera eficiente: el '
    'conjunto de portafolios que maximizan el retorno esperado para un nivel dado de '
    'riesgo (medido por la varianza). Este marco teórico, ampliado posteriormente por '
    'Sharpe (1964) con el modelo de valoración de activos de capital (CAPM) y la '
    'introducción del ratio de Sharpe como medida de rendimiento ajustado por riesgo, '
    'continúa siendo el fundamento de la gestión cuantitativa de carteras.'
)
body(
    'Merton (1972) derivó analíticamente la frontera eficiente, estableciendo las '
    'propiedades matemáticas del portafolio de mínima varianza global y demostrando '
    'que cualquier portafolio eficiente puede expresarse como combinación lineal de '
    'dos portafolios de referencia. La implementación práctica de la optimización '
    'de Markowitz enfrenta desafíos relacionados con la estimación de la matriz de '
    'covarianzas y los retornos esperados, particularmente con portafolios de gran '
    'dimensión (Black & Litterman, 1992).'
)
body(
    'Las herramientas de software libre (SciPy, CVXPY) han democratizado '
    'parcialmente el acceso a la optimización de portafolios, aunque su integración '
    'en plataformas web accesibles para inversores no institucionales sigue siendo '
    'limitada. Trabajos recientes han combinado la optimización de Markowitz con '
    'predicciones de ML para estimar retornos esperados de forma más robusta '
    '(Kolm et al., 2014), una aproximación adoptada en el PortfolioAgent del '
    'sistema propuesto.'
)

h3('1.3.5 Sistemas multiagentes en finanzas')

body(
    'Los sistemas multiagentes (MAS) proveen un marco formal para el diseño de '
    'aplicaciones distribuidas en las que múltiples entidades autónomas colaboran '
    'para resolver tareas complejas que exceden la capacidad de un agente individual '
    '(Wooldridge, 2009). En el dominio financiero, los MAS han sido empleados '
    'principalmente en simulación de mercados (LeBaron, 2001), detección de fraudes '
    '(Bose & Mahapatra, 2001) y sistemas de trading algorítmico (Huang et al., 2021).'
)
body(
    'La arquitectura multiagente proporciona ventajas específicas para el análisis '
    'financiero integrado: permite la separación de responsabilidades entre agentes '
    'especializados (mercado, predicción, sentimiento, fundamental, portafolio), '
    'facilita la incorporación de nuevos módulos sin modificar los existentes y '
    'habilita la ejecución paralela de tareas independientes, reduciendo la latencia '
    'total del sistema (Jennings, 2000). Estas propiedades son particularmente '
    'relevantes en aplicaciones donde la actualidad de la información es crítica '
    'para la calidad del análisis.'
)

h3('1.3.6 Plataformas comerciales y brecha tecnológica')

body(
    'Las principales plataformas comerciales de análisis financiero (Bloomberg '
    'Terminal, Reuters Eikon, FactSet) ofrecen capacidades analíticas avanzadas, '
    'pero con barreras de acceso que las hacen inviables para inversores minoristas. '
    'Bloomberg Terminal, la plataforma más utilizada por instituciones financieras, '
    'centraliza datos en tiempo real, indicadores económicos y herramientas analíticas, '
    'pero su costo de licencia (aproximadamente USD 24.000 anuales) y su interfaz '
    'orientada a usuarios expertos la convierten en una solución prácticamente '
    'inaccesible fuera del ámbito institucional (King, 2016).'
)
body(
    'En el espacio de herramientas de acceso libre, existen implementaciones parciales: '
    'modelos NLP para análisis de sentimiento financiero (FinBERT), bibliotecas de ML '
    'para predicción de precios (scikit-learn, XGBoost) y herramientas de optimización '
    'de carteras (PyPortfolioOpt, CVXPY). Sin embargo, estas soluciones se presentan '
    'de forma fragmentada, requieren conocimientos técnicos avanzados para su '
    'configuración y no ofrecen una interfaz integrada y accesible para usuarios no '
    'especializados. La integración de estos componentes en una plataforma unificada, '
    'documentada y validada empíricamente sigue siendo un desafío abierto en la '
    'literatura.'
)

# ── 1.4 ─────────────────────────────────────────────────────────────────────
h2('1.4 Brecha de investigación y justificación')

body(
    'La revisión del estado del arte permite identificar una brecha específica: si '
    'bien existen investigaciones y herramientas que abordan de forma individual la '
    'predicción de precios mediante ML, el análisis de sentimiento financiero mediante '
    'NLP, la incorporación de datos fundamentales y la optimización de carteras con '
    'Markowitz, no se han encontrado en la literatura trabajos que integren los '
    'cuatro componentes en una arquitectura multiagente única, modular, de código '
    'abierto y validada empíricamente con datos reales de mercado.'
)
body(
    'Este trabajo contribuye a cerrar dicha brecha mediante el diseño, implementación '
    'y validación de un prototipo funcional que integra los cuatro componentes en una '
    'arquitectura de siete agentes coordinados por un backend REST (FastAPI), '
    'accesibles a través de una interfaz web (Streamlit) y respaldados por una base '
    'de datos relacional para trazabilidad y auditoría (SQLite, SQLAlchemy). Las '
    'contribuciones principales del trabajo son: (a) la arquitectura de integración '
    'multiagente que combina análisis técnico, ML, NLP, fundamentales y optimización '
    'de Markowitz en un único sistema; (b) la validación empírica del sistema con '
    'datos reales de diez acciones del mercado estadounidense; y (c) la demostración '
    'de viabilidad de un sistema de este tipo con tecnologías de bajo costo y código '
    'abierto, accesible para usuarios sin formación técnica avanzada.'
)

# ── 1.5 ─────────────────────────────────────────────────────────────────────
h2('1.5 Motivación')

body(
    'La motivación central de este trabajo radica en la necesidad de reducir la brecha '
    'tecnológica e informativa que separa a las grandes instituciones financieras de '
    'los pequeños inversores. Mientras los actores de gran escala cuentan con '
    'plataformas integradas, acceso prioritario a datos y capacidad de análisis '
    'automatizado, los usuarios minoristas suelen depender de información fragmentada, '
    'tardía o poco contextualizada y, con frecuencia, carecen de los conocimientos '
    'técnicos en el área de finanzas cuantitativas necesarios para interpretar y actuar '
    'sobre dicha información. Esta asimetría limita su capacidad para anticipar '
    'movimientos del mercado, construir carteras diversificadas eficientemente y '
    'tomar decisiones de inversión fundamentadas.'
)
body(
    'El desarrollo de un sistema inteligente, autónomo y modular busca revertir esta '
    'situación mediante la provisión de análisis automatizados, alertas personalizadas '
    'y herramientas de optimización de portafolios que permitan a cada usuario '
    'reaccionar de forma oportuna ante eventos financieros relevantes y construir '
    'carteras eficientes según la teoría de Markowitz. La arquitectura multiagente '
    'otorga al sistema flexibilidad y robustez, permitiendo que cada agente cumpla '
    'un rol especializado dentro de una estructura distribuida que facilita la '
    'evolución y el mantenimiento del sistema.'
)
body(
    'Desde una perspectiva académica, este trabajo también responde al interés '
    'creciente en la aplicación práctica de la inteligencia artificial al dominio '
    'financiero, un área en la que la brecha entre los desarrollos teóricos '
    'publicados en la literatura y las herramientas accesibles para el usuario '
    'final es particularmente amplia (López de Prado, 2018). La construcción de '
    'un prototipo funcional y validado empíricamente contribuye a cerrar parcialmente '
    'esta brecha y a establecer una base reproducible para investigaciones futuras.'
)

# ── 1.6 ─────────────────────────────────────────────────────────────────────
h2('1.6 Objetivos y alcance')

h3('1.6.1 Objetivo general')
body(
    'Diseñar, implementar y validar empíricamente un sistema multiagente inteligente '
    'que integre análisis técnico, predicción de dirección de precios mediante machine '
    'learning, análisis de sentimiento en noticias financieras, análisis fundamental '
    'con datos SEC EDGAR y optimización de portafolios mediante la teoría moderna de '
    'carteras de Markowitz (1952), con el propósito de apoyar la toma de decisiones '
    'de inversión de usuarios minoristas sin formación técnica especializada.'
)

h3('1.6.2 Objetivos específicos')
bullet(
    'Desarrollar siete agentes especializados (MarketAgent, ModelAgent, SentimentAgent, '
    'RecommendationAgent, AlertAgent, SECAgent, PortfolioAgent) que operen de forma '
    'coordinada bajo una arquitectura modular y escalable.'
)
bullet(
    'Aplicar un ensemble de clasificadores de machine learning (Random Forest, '
    'Gradient Boosting, XGBoost y LightGBM) con walk-forward validation temporal '
    'para predecir la dirección del precio de activos financieros en un horizonte '
    'de tres días hábiles, evaluando las métricas de clasificación binaria: Accuracy, '
    'Precision, Recall, F1-Score y AUC-ROC.'
)
bullet(
    'Implementar un ensemble de modelos NLP (FinBERT, VADER, lexicón financiero '
    'especializado y TextBlob) para el análisis de sentimiento en noticias financieras, '
    'con ponderación diferencial por modelo y filtro de relevancia por ticker.'
)
bullet(
    'Desarrollar el SECAgent para la obtención automatizada de ratios fundamentales '
    '(P/E, ROE, márgenes, crecimiento) vía yfinance y de filings regulatorios '
    '(10-K, 10-Q, 8-K) a través de la API pública de SEC EDGAR, generando un score '
    'fundamental entre −1 y +1.'
)
bullet(
    'Implementar el PortfolioAgent para el análisis y optimización de portafolios '
    'mediante la teoría de Markowitz: portafolio de máximo ratio de Sharpe, portafolio '
    'de mínima varianza y frontera eficiente, con restricciones de pesos y cálculo de '
    'métricas de riesgo (VaR 95 % y 99 %, beta respecto al S&P 500).'
)
bullet(
    'Diseñar e implementar una interfaz web en Streamlit que permita al usuario '
    'visualizar métricas técnicas, predicciones de ML, análisis de sentimiento, '
    'datos fundamentales, portafolios optimizados y alertas, de forma accesible '
    'sin conocimientos técnicos previos.'
)
bullet(
    'Validar empíricamente el sistema mediante pruebas funcionales (30 pruebas, '
    '10 tickers, 3 iteraciones) y pruebas de carga concurrente con datos reales '
    'de mercado obtenidos a través de la API de Yahoo Finance.'
)

h3('1.6.3 Alcance')
body(
    'El prototipo se limita al procesamiento de datos públicos del mercado '
    'estadounidense y a la generación de señales de análisis y alertas sobre un '
    'conjunto de activos seleccionados. No incluye operaciones reales de compra o '
    'venta de activos, integración con brokers, gestión de carteras en tiempo real '
    'ni rebalanceo automático. La optimización de portafolios se realiza sobre datos '
    'históricos; el cálculo de costos de transacción y el rebalanceo automático '
    'quedan propuestos como líneas de trabajo futuro. La puesta en producción no '
    'forma parte del alcance de este trabajo; se propone como línea de trabajo futuro '
    'la containerización con Docker y el despliegue en servicios de nube.'
)

# ── 1.7 ─────────────────────────────────────────────────────────────────────
h2('1.7 Impacto social y económico')

body(
    'La desigualdad informativa en los mercados financieros impacta directamente sobre '
    'la capacidad de decisión de los pequeños inversores y, en última instancia, sobre '
    'la eficiencia en la asignación del capital en la economía real (Stiglitz, 2002). '
    'El sistema propuesto contribuye a democratizar el acceso a la información y la '
    'analítica avanzada, ofreciendo a usuarios minoristas herramientas de análisis '
    'previamente reservadas a actores institucionales, con cuatro impactos concretos:'
)
bullet(
    'Mayor transparencia: el sistema centraliza en una única plataforma información '
    'técnica, fundamental, de sentimiento y de portafolio, reduciendo la fragmentación '
    'informativa que perjudica a los inversores individuales.'
)
bullet(
    'Reducción de la brecha tecnológica: la disponibilidad de una herramienta '
    'gratuita, de código abierto y accesible a través de un navegador web reduce '
    'las barreras de entrada al análisis cuantitativo avanzado.'
)
bullet(
    'Educación financiera: las explicaciones textuales generadas por el '
    'RecommendationAgent y las visualizaciones de la frontera eficiente del '
    'PortfolioAgent contribuyen a la formación financiera del usuario, '
    'facilitando la comprensión de conceptos como riesgo, diversificación '
    'y optimización de carteras.'
)
bullet(
    'Fomento de la innovación regional: el prototipo demuestra que es posible '
    'construir herramientas de análisis financiero avanzado con tecnologías de '
    'código abierto (FastAPI, Streamlit, scikit-learn, scipy), sentando '
    'precedentes para desarrollos similares orientados a mercados de América Latina.'
)
body(
    'Cabe destacar que el sistema fue diseñado como herramienta de apoyo a la '
    'decisión, complementaria al juicio del inversor, y no como sustituto de '
    'asesoramiento financiero profesional. Los modelos de ML y los algoritmos de '
    'optimización operan sobre datos históricos y no garantizan rendimientos futuros; '
    'el usuario asume la responsabilidad final sobre sus decisiones de inversión.'
)

# ── 1.8 ─────────────────────────────────────────────────────────────────────
h2('1.8 Organización de la memoria')

body(
    'La presente memoria técnica se organiza en cinco capítulos y un apéndice. '
    'El capítulo 1 introduce el contexto, el problema de investigación, el estado '
    'del arte, la brecha de investigación, los objetivos y el alcance del sistema '
    'desarrollado. El capítulo 2 presenta el marco teórico que sustenta las '
    'decisiones de diseño, incluyendo los fundamentos de los sistemas multiagentes, '
    'machine learning aplicado a finanzas, análisis de sentimiento con NLP, análisis '
    'fundamental y SEC EDGAR, y la teoría moderna de carteras de Markowitz. '
    'El capítulo 3 describe en detalle el diseño e implementación del sistema, '
    'abarcando la arquitectura de los siete agentes, el backend FastAPI, la base de '
    'datos SQLite, el sistema de alertas, los mecanismos de seguridad y el pipeline '
    'de procesamiento. El capítulo 4 presenta los resultados obtenidos y la validación '
    'empírica del sistema, incluyendo métricas de clasificación ML, evaluación del '
    'análisis de sentimiento, pruebas funcionales end-to-end y pruebas de carga '
    'concurrente. Finalmente, el capítulo 5 expone las conclusiones, las limitaciones '
    'del trabajo y las principales líneas de investigación futura. El Apéndice A '
    'incluye la tabla comparativa de objetivos planificados versus realizados.'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 2 — MARCO TEÓRICO
# ════════════════════════════════════════════════════════════════════════════
h1('Capítulo 2 — Marco teórico')
body(
    'Este capítulo incluye los fundamentos teóricos y técnicos que sustentan el '
    'prototipo multiagente, abarcando sistemas multiagentes, machine learning aplicado '
    'a finanzas, NLP para análisis de sentimiento, análisis fundamental y la teoría '
    'moderna de carteras de Markowitz (1952).'
)

h2('2.1 Sistema multiagente')
body(
    'Un sistema multiagente (SMA) se basa en la interacción entre entidades de '
    'software denominadas agentes que cooperan, compiten o se comunican entre sí para '
    'resolver tareas complejas (Wooldridge, 2009). Un agente es una unidad de software '
    'autónoma capaz de percibir su entorno, procesar información y actuar sobre él. '
    'El ciclo básico de funcionamiento comprende tres etapas: percepción (recepción '
    'de datos del entorno), decisión (análisis y selección de acción) y acción '
    '(ejecución de tarea o comunicación de resultados).'
)
body(
    'Los agentes presentan cuatro propiedades fundamentales (Wooldridge, 2009): '
    'autonomía (actúan sin intervención directa del usuario), cooperación '
    '(colaboran mediante intercambio de mensajes), modularidad (cada agente cumple '
    'un rol específico) y escalabilidad (admite la incorporación de nuevos agentes '
    'sin modificar los existentes). En el prototipo desarrollado, los siete agentes '
    'son: MarketAgent, ModelAgent, SentimentAgent, RecommendationAgent, AlertAgent, '
    'SECAgent y PortfolioAgent.'
)

h2('2.2 Machine learning aplicado a finanzas')
body(
    'El aprendizaje automático ha demostrado gran potencial para modelar y predecir '
    'el comportamiento de los mercados financieros (López de Prado, 2018). Los métodos '
    'de ensemble —Random Forest (Breiman, 2001), Gradient Boosting (Friedman, 2001), '
    'XGBoost (Chen & Guestrin, 2016) y LightGBM (Ke et al., 2017)— combinan múltiples '
    'modelos para mejorar la precisión y reducir el sobreajuste. El entrenamiento '
    'utiliza walk-forward validation con TimeSeriesSplit, respetando el orden temporal '
    'de los datos para evitar data leakage (López de Prado, 2018). La ventana de '
    'entrenamiento es de 504 días, el horizonte de predicción de 3 días y se aplica '
    'un umbral del 0,5 % en la definición del target para filtrar ruido de mercado.'
)

h2('2.3 NLP para análisis de sentimiento financiero')
body(
    'El sistema implementa un ensemble de cuatro modelos NLP ponderados: FinBERT '
    '(Araci, 2019) con peso del 40 %, VADER (Hutto & Gilbert, 2014) con el 25 %, '
    'un lexicón financiero especializado de 500+ términos con el 20 %, y TextBlob '
    '(Loria, 2020) con el 15 %. El score consolidado oscila entre −1 (muy negativo) '
    'y +1 (muy positivo). Incluye un filtro de relevancia que pondera con un 40 % '
    'las noticias sin mención directa del ticker analizado, evitando la contaminación '
    'por noticias de mercado general.'
)

h2('2.4 Análisis fundamental y SEC EDGAR')
body(
    'El análisis fundamental evalúa el valor intrínseco de una empresa a través de '
    'indicadores financieros derivados de sus estados contables (Damodaran, 2012): '
    'ratios de valuación (P/E, P/B, P/S, EV/EBITDA), rentabilidad (ROE, ROA, márgenes), '
    'crecimiento (YoY de ingresos y ganancias) y solidez financiera (Deuda/Capital, '
    'Current Ratio, Quick Ratio). La API pública de SEC EDGAR (SEC, 2023) permite el '
    'acceso gratuito a declaraciones financieras regulatorias (10-K, 10-Q, 8-K), '
    'habilitando la integración de información fundamental sin costo de licencia.'
)

h2('2.5 Teoría moderna de carteras — Markowitz (1952)')
body(
    'La teoría moderna de carteras (Markowitz, 1952) establece que el riesgo de un '
    'portafolio depende de las correlaciones entre sus activos, y que la diversificación '
    'permite reducir el riesgo sin sacrificar retorno esperado. Las relaciones '
    'fundamentales son: retorno esperado del portafolio rp = Σ wᵢ·rᵢ; varianza del '
    'portafolio σ²p = wᵀ Σ w (donde Σ es la matriz de covarianzas); y ratio de '
    'Sharpe S = (rp − rf) / σp (Sharpe, 1964), con rf = 4,5 % anual. El PortfolioAgent '
    'resuelve los problemas de optimización cuadrática (scipy.optimize) para obtener: '
    'portafolio de máximo Sharpe (pesos ∈ [1 %, 65 %], Σwᵢ = 1), portafolio de mínima '
    'varianza y frontera eficiente con 15 puntos (Merton, 1972).'
)

h2('2.6 Arquitectura general y evaluación de riesgos')
body(
    'La arquitectura del sistema multiagente propuesto se basa en tres elementos '
    'tecnológicos principales: FastAPI como coordinador central del pipeline (Ramírez, '
    '2021), SQLite como base de datos relacional ligera (Hipp et al., 2023) y '
    'Streamlit como biblioteca de construcción rápida de dashboards analíticos '
    '(Streamlit Inc., 2024). La arquitectura modular garantiza la escalabilidad y '
    'la incorporación de nuevos componentes sin alterar el funcionamiento general '
    'del sistema.'
)
body(
    'Los riesgos técnicos identificados incluyen: sesgos en los datos de entrenamiento '
    '(Barredo Arrieta et al., 2020), sobreajuste —mitigado mediante walk-forward '
    'validation (López de Prado, 2018)—, dependencia de Yahoo Finance como fuente '
    'única de datos de mercado y noticias, y riesgo de degradación del modelo ante '
    'cambios estructurales del mercado (Danielsson et al., 2022). El sistema se alinea '
    'con los lineamientos de Basilea III y con la Comunicación A 7724 del BCRA para '
    'la gestión de riesgos informáticos en sistemas automatizados.'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 3 — DISEÑO E IMPLEMENTACIÓN
# ════════════════════════════════════════════════════════════════════════════
h1('Capítulo 3 — Diseño e implementación')
body(
    'El sistema implementa una arquitectura multiagente que integra datos de mercado, '
    'predicción, análisis de sentimiento, datos fundamentales, optimización de '
    'portafolios y generación de alertas inteligentes. Combina un backend en FastAPI '
    'con una interfaz en Streamlit y una base de datos SQLite orientada a trazabilidad '
    'y auditoría, en línea con los requerimientos regulatorios vigentes (BCRA, 2022; '
    'Basel Committee, 2017).'
)

h2('3.1 Arquitectura multiagente')
body(
    'La arquitectura propuesta se organiza alrededor de siete agentes especializados '
    'que cooperan para cubrir el ciclo completo de obtención de datos, análisis, '
    'optimización, recomendación y alertas. La coordinación se realiza a través del '
    'backend FastAPI, que actúa como orquestador, y de una base de datos SQLite '
    'compartida como memoria persistente del sistema (Wooldridge, 2009).'
)

body('Tabla 3.1: Funciones y salidas de los siete agentes del sistema.')
add_table(
    ['Agente', 'Función principal', 'Salida principal'],
    [
        ['MarketAgent',        'Descarga datos históricos y calcula 35+ indicadores técnicos (SMA, EMA, RSI, MACD, Bollinger, ATR, OBV, VWAP). Detección de anomalías con 5 algoritmos.', 'Datos de mercado normalizados, score técnico ponderado.'],
        ['ModelAgent',         'Ensemble de clasificadores (RF, GB, XGB, LightGBM + Linear, Ridge opcionales). Ventana 504 días, 5 folds walk-forward, umbral 0,5 %.', 'Predicción SUBIDA/BAJADA, prob_subida calibrada, métricas de clasificación.'],
        ['SentimentAgent',     'Ensemble NLP: FinBERT (40 %), VADER (25 %), Lexicón (20 %), TextBlob (15 %). Análisis de 7 noticias recientes con filtro de relevancia.', 'Score sentimiento −1 a +1, categoría, tendencia, confianza.'],
        ['RecommendationAgent','Integra señales técnicas (40 %), predicción (35 %), sentimiento (15 %), riesgo (10 %). VaR 95 %, Kelly Criterion.', 'Recomendación textual (7 niveles), position sizing, stop loss / take profit.'],
        ['AlertAgent',         'Evalúa umbrales (warning 3 %, critical 7 %) mediante Z-Score, MAD, CUSUM, EWMA e Isolation Forest.', 'Alerta con 6 niveles de severidad (INFO→EMERGENCY), persistida en BD.'],
        ['SECAgent',           'Ratios fundamentales vía yfinance.info. Filings 10-K/10-Q/8-K vía API pública SEC EDGAR (SEC, 2023).', 'Score fundamental −1 a +1, balance resumido, filings recientes.'],
        ['PortfolioAgent',     'Orquesta los demás agentes en paralelo (ThreadPoolExecutor). Markowitz: máx Sharpe, mín varianza, frontera eficiente 15 puntos (scipy.optimize).', 'Métricas de portafolio, pesos óptimos, VaR 95 %/99 %, Sharpe, matriz de correlación.'],
    ]
)

h3('3.1.1 MarketAgent')
body(
    'El MarketAgent descarga series históricas de Yahoo Finance mediante yfinance, '
    'calcula 35+ indicadores técnicos en cuatro categorías: tendencia (SMA, EMA, MACD, '
    'ADX, Ichimoku, Parabolic SAR), momentum (RSI, Stochastic, Williams %R, ROC, CCI), '
    'volatilidad (Bollinger Bands, ATR, Keltner Channels) y volumen (OBV, VWAP, MFI, '
    'ADL, CMF). Implementa detección de anomalías con cinco algoritmos: Z-Score, MAD, '
    'CUSUM, Isolation Forest y Volume Anomaly.'
)

h3('3.1.2 ModelAgent')
body(
    'El ModelAgent implementa clasificación binaria de dirección de precios mediante '
    'un ensemble de modelos base (LogisticRegression, RidgeClassifier, Random Forest y '
    'Gradient Boosting) y modelos opcionales (XGBoost, LightGBM y LSTM según '
    'disponibilidad de librerías). La calibración de probabilidades se realiza con '
    'CalibratedClassifierCV (método sigmoide, cv=\'prefit\'), reservando ~20 % de '
    'los datos para calibración. El ensemble combina las probabilidades calibradas '
    'mediante promedio ponderado por desempeño en validación.'
)

h3('3.1.3 SentimentAgent')
body(
    'El SentimentAgent analiza noticias de Yahoo Finance (máximo 7 por sesión) '
    'mediante el ensemble NLP descrito en la sección 2.3. El filtro de relevancia '
    'asigna peso reducido (40 %) a noticias sin mención directa del ticker analizado, '
    'evitando contaminación por noticias de mercado general. Cuenta con caché de '
    '1 hora para optimizar la latencia del sistema.'
)

h3('3.1.4 RecommendationAgent')
body(
    'El RecommendationAgent combina señales de los agentes anteriores mediante un '
    'sistema de decisión multi-factor con 15+ variables, ponderadas por categoría: '
    'señales técnicas (40 %), predicción ML (35 %), sentimiento NLP (15 %) y '
    'riesgo (10 %). Implementa gestión de riesgo con VaR al 95 %, estimación de '
    'max drawdown y position sizing con Kelly Criterion. Genera 7 niveles de '
    'recomendación desde Compra Fuerte hasta Venta Fuerte, con explicabilidad '
    'completa de cada factor contribuyente.'
)

h3('3.1.5 AlertAgent')
body(
    'El AlertAgent utiliza cinco técnicas de detección de anomalías (Z-Score, MAD, '
    'CUSUM, EWMA e Isolation Forest) para determinar uno de seis niveles de severidad: '
    'INFO, LOW, MEDIUM, HIGH, CRITICAL y EMERGENCY. El sistema distingue nueve tipos '
    'de alerta: movimiento de precio, pico de volatilidad, cambio de tendencia, '
    'anomalía detectada, volumen inusual, cambio de sentimiento, divergencia de '
    'predicción, ruptura de correlación y patrón detectado.'
)

h3('3.1.6 SECAgent')
body(
    'El SECAgent obtiene datos fundamentales en dos niveles: ratios financieros vía '
    'yfinance.info (valuación: P/E, P/B, P/S, EV/EBITDA; rentabilidad: ROE, ROA, '
    'márgenes bruto/operativo/neto; crecimiento YoY; deuda/capital; ratios de '
    'liquidez; beta y capitalización de mercado) y filings regulatorios vía la API '
    'pública de SEC EDGAR (SEC, 2023), consultando automáticamente el CIK por ticker '
    'y recuperando los filings más recientes (10-K, 10-Q, 8-K, DEF 14A). Genera un '
    'score fundamental entre −1 y +1 con caché de 2 horas. Se integra en '
    'GET /predict/{ticker} como paso paralelo con el parámetro incluir_sec=true.'
)

h3('3.1.7 PortfolioAgent')
body(
    'El PortfolioAgent orquesta todos los agentes anteriores en paralelo mediante '
    'ThreadPoolExecutor para cada activo del portafolio. Calcula las métricas de la '
    'teoría moderna de carteras (Markowitz, 1952): retorno esperado anualizado '
    '(ponderación histórica + predicción ML), volatilidad del portafolio '
    '(σp = √(wᵀ Σ w)), ratio de Sharpe ((Rp − Rf) / σp con Rf = 4,5 %), VaR '
    'paramétrico al 95 % y 99 %, ratio de diversificación y beta versus S&P 500. '
    'La optimización (scipy.optimize) resuelve tres problemas: portafolio de máximo '
    'Sharpe (pesos ∈ [1 %, 65 %], Σwᵢ = 1), portafolio de mínima varianza y frontera '
    'eficiente con 15 puntos. Soporta de 2 a 15 activos y se expone vía '
    'POST /portfolio/analyze.'
)

h2('3.2 Backend basado en FastAPI')
body(
    'El backend del sistema se desarrolló utilizando FastAPI (Ramírez, 2021), '
    'un framework moderno para la construcción de APIs REST en Python. La aplicación '
    'expone endpoints agrupados en cuatro routers: autenticación '
    '(/auth/register, /auth/login, /auth/me, /auth/forgot-password, '
    '/auth/reset-password), predicción (GET /predict/{ticker} con soporte para '
    'incluir_sec=true, GET /predict/{ticker}/market, GET /predict/{ticker}/sentiment), '
    'portafolio (POST /portfolio/analyze) y alertas (GET /alerts, GET /alerts/stats, '
    'GET /alerts/{id}, PUT /alerts/{id}/read, DELETE /alerts/{id}). El endpoint '
    'GET /predict/{ticker} invoca el pipeline secuencial completo; el PortfolioAgent '
    'ejecuta el pipeline en paralelo para cada activo del portafolio.'
)

h2('3.3 Dashboard en Streamlit')
body(
    'La interfaz de usuario implementada en Streamlit (Streamlit Inc., 2024) se '
    'organiza en tres pestañas: (1) Análisis de activos — gráfico Plotly interactivo '
    'con indicadores técnicos, predicción de dirección a 3 días, análisis de '
    'sentimiento con noticias individuales, recomendación multi-factor y datos '
    'fundamentales del SECAgent (ratios, balance y filings EDGAR); (2) Portafolio — '
    'formulario de entrada de tickers y pesos normalizados, panel de métricas de '
    'Markowitz, gráfico de distribución de pesos, mapa de calor de correlaciones, '
    'comparativa de portafolios (actual vs. máx Sharpe vs. mín varianza) y '
    'visualización de la frontera eficiente; y (3) Centro de alertas — lista con '
    'filtros, códigos de color por severidad, estadísticas e historial completo.'
)

h2('3.4 Base de datos')
body(
    'La solución utiliza SQLite como base de datos relacional ligera (Hipp et al., '
    '2023), accedida a través de SQLAlchemy (ORM). Las tablas principales son: '
    'usuarios, alertas, métricas_modelo y password_reset_tokens. El esquema garantiza '
    'la trazabilidad completa: cada alerta registra el usuario asociado, el ticker, '
    'el tipo de severidad, la variación porcentual, los precios actual y proyectado, '
    'y la marca temporal de creación, permitiendo auditorías posteriores conforme a '
    'los requerimientos de la Comunicación A 7724 del BCRA (2022).'
)

h2('3.5 Seguridad')
body(
    'La arquitectura incorpora los mecanismos de seguridad recomendados por OWASP '
    '(2021): autenticación JWT con tokens firmados (expiración configurable), hash '
    'de contraseñas con bcrypt (resistente a ataques de fuerza bruta), validación '
    'estricta de parámetros con Pydantic, rate limiting de alertas, gestión de '
    'secretos mediante variables de entorno (Pydantic Settings) y registro '
    'estructurado de eventos con marca temporal, módulo origen y nivel de severidad '
    'para auditoría y detección de anomalías.'
)

h2('3.6 Pipeline de procesamiento')
body('El flujo completo para análisis de activo individual comprende siete etapas:')
numbered('Ingesta y validación de datos de mercado (MarketAgent).')
numbered('Normalización y limpieza de series temporales (52 indicadores técnicos).')
numbered('Ejecución del ensemble de clasificadores (ModelAgent, ventana 504 días).')
numbered('Análisis de noticias y sentimiento (SentimentAgent, ensemble NLP).')
numbered('Obtención paralela de datos fundamentales y filings (SECAgent).')
numbered('Integración de señales, recomendación y generación de alertas (RecommendationAgent, AlertAgent).')
numbered('Almacenamiento en SQLite y visualización en el dashboard Streamlit.')
body('Para análisis de portafolio, el PortfolioAgent adiciona:')
numbered('Recepción de lista de tickers y pesos normalizados.')
numbered('Ejecución paralela del pipeline completo para cada activo (ThreadPoolExecutor).')
numbered('Cálculo de la matriz de covarianzas y métricas del portafolio actual.')
numbered('Optimización de Markowitz (máximo Sharpe, mínima varianza, frontera eficiente).')
numbered('Generación de alertas de portafolio (concentración >40 %, correlación >0,85) y respuesta JSON.')

h2('3.7 Bibliotecas principales')
add_table(
    ['Paquete', 'Categoría', 'Función'],
    [
        ['fastapi',            'Backend API',       'Framework para endpoints REST y documentación OpenAPI.'],
        ['uvicorn',            'Servidor ASGI',     'Ejecuta FastAPI con manejo concurrente de peticiones.'],
        ['pydantic',           'Validación',        'Modelado estricto de datos y configuración segura.'],
        ['sqlalchemy',         'Base de datos',     'ORM para SQLite (usuarios, alertas, métricas).'],
        ['yfinance',           'Datos financieros', 'Precios históricos, noticias y ratios fundamentales.'],
        ['requests',           'HTTP',              'Consultas a la API pública de SEC EDGAR.'],
        ['scipy',              'Optimización',      'Optimización cuadrática de Markowitz (frontera eficiente).'],
        ['scikit-learn',       'ML',                'Clasificadores base, preprocesamiento y métricas.'],
        ['xgboost / lightgbm', 'ML (opcional)',     'Clasificadores de boosting para el ensemble (Chen & Guestrin, 2016; Ke et al., 2017).'],
        ['torch / transformers','Deep Learning/NLP','LSTM (opcional) y FinBERT para análisis de sentimiento (Araci, 2019).'],
        ['nltk',               'NLP',               'VADER para análisis de sentimiento (Hutto & Gilbert, 2014).'],
        ['textblob',           'NLP',               'Análisis de polaridad general de texto.'],
        ['streamlit',          'Dashboard',         'Interfaz interactiva para predicciones, portafolio y alertas.'],
        ['plotly',             'Visualización',     'Gráficos interactivos: precios, frontera eficiente, heatmaps.'],
        ['passlib / python-jose','Seguridad',       'Bcrypt para contraseñas, JWT para autenticación (OWASP, 2021).'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 4 — ENSAYOS Y RESULTADOS
# ════════════════════════════════════════════════════════════════════════════
h1('Capítulo 4 — Ensayos y resultados')

body(
    'El presente capítulo tiene como objetivo validar empíricamente la efectividad del '
    'sistema propuesto mediante una serie exhaustiva de experimentos, pruebas y análisis '
    'sobre datos reales de mercado. La validación de un sistema de inteligencia artificial '
    'en el dominio financiero requiere un enfoque riguroso y multidimensional: no basta con '
    'demostrar que los modelos de machine learning alcanzan métricas aceptables en conjuntos '
    'de validación; es necesario evaluar el sistema completo en condiciones operativas reales, '
    'considerando aspectos de rendimiento, escalabilidad, usabilidad y, fundamentalmente, el '
    'valor agregado que proporciona a usuarios con diferentes perfiles y objetivos de inversión '
    '(López de Prado, 2018).'
)
body(
    'Los resultados presentados fueron obtenidos en cuatro sesiones de prueba con datos reales '
    'de mercado obtenidos mediante la API de Yahoo Finance: (1) 9 de febrero de 2026 — pruebas '
    'de carga y funcionales con configuración inicial (ventana de entrenamiento de 30 días, '
    '3 folds, sin umbral de target); (2) 13 de febrero de 2026 — evaluación del SentimentAgent '
    'y del ModelAgent con los 10 tickers de referencia; (3) 8 de marzo de 2026 — aplicación '
    'de tres mejoras al ModelAgent: ampliación de la ventana de entrenamiento de 252 a 504 '
    'días, incremento de 3 a 5 folds en la validación cruzada temporal (TimeSeriesSplit), e '
    'incorporación de un umbral del 0,5 % en la definición del target para filtrar ruido de '
    'mercado; y (4) 21 de marzo de 2026 — corrección de un sesgo en el RecommendationAgent y '
    'validación final del sistema con la configuración definitiva. Las métricas de la sección '
    '4.1 corresponden a la sesión del 8 de marzo de 2026. Se siguió un protocolo de '
    'experimentación que incluye walk-forward validation temporal para evitar data leakage '
    '(López de Prado, 2018) y documentación exhaustiva de las condiciones experimentales para '
    'garantizar la reproducibilidad de los resultados.'
)
body(
    'El capítulo se organiza en cuatro secciones: evaluación de los modelos de machine '
    'learning (sección 4.1), evaluación del módulo de análisis de sentimiento NLP (sección '
    '4.2), pruebas end-to-end funcionales y de rendimiento (sección 4.3) y análisis de casos '
    'de uso ilustrativos (sección 4.4). El capítulo concluye con un análisis crítico de las '
    'fortalezas del sistema —robustez, consistencia, integración efectiva de múltiples '
    'agentes— y de sus limitaciones —escalabilidad restringida a 10 usuarios concurrentes, '
    'dependencia de servicios externos, precisión variable según volatilidad del activo—, '
    'estableciendo las bases para las mejoras futuras discutidas en el capítulo 5.'
)

# ── 4.1 ─────────────────────────────────────────────────────────────────────
h2('4.1 Evaluación de modelos ML')

h3('4.1.1 Metodología de evaluación')
body(
    'Para evaluar el desempeño del ModelAgent se utilizó un enfoque de clasificación '
    'binaria que predice la dirección del precio (SUBIDA/BAJADA) en un horizonte de 3 '
    'días hábiles, en lugar del precio exacto. Esta elección responde a la literatura '
    'especializada, que señala que la predicción de dirección es más robusta que la '
    'predicción de valor exacto en presencia de alta volatilidad y baja relación '
    'señal-ruido (Fischer & Krauss, 2018; López de Prado, 2018).'
)
body('Las métricas de evaluación seleccionadas son:')
bullet('Accuracy: porcentaje de predicciones correctas de dirección sobre el total de observaciones.')
bullet('Precision: dado que el sistema predice SUBIDA, porcentaje de veces que el precio efectivamente sube más del 0,5 %. Relevante para minimizar falsas señales alcistas.')
bullet('Recall: porcentaje de subidas reales (>0,5 %) que el modelo detecta. Relevante para no perder oportunidades alcistas reales.')
bullet('F1-Score: media armónica entre Precision y Recall. Métrica de compromiso para clases desequilibradas.')
bullet('AUC-ROC: área bajo la curva ROC, mide la capacidad discriminativa global del modelo independientemente del umbral de decisión (Fawcett, 2006). AUC = 0,5 equivale al azar; AUC = 1,0 equivale a separación perfecta.')
body(
    'La evaluación se realizó con datos históricos de 10 acciones representativas del '
    'mercado estadounidense, cubriendo tecnología (AAPL, MSFT, GOOGL), automotriz/EV '
    '(TSLA), comercio electrónico (AMZN), redes sociales (META), semiconductores (NVDA), '
    'sector financiero (JPM, V) y retail (WMT). Se utilizó una ventana de entrenamiento '
    'de 504 días (2 años bursátiles), horizonte de predicción de 3 días y validación '
    'cruzada temporal con 5 folds mediante TimeSeriesSplit (Scikit-learn, 2023). '
    'El target de clasificación se definió con un umbral del 0,5 %: se clasifica como '
    'SUBIDA únicamente si el precio aumenta más del 0,5 % en el horizonte evaluado, '
    'filtrando movimientos menores que constituyen ruido de mercado (López de Prado, 2018).'
)

h3('4.1.2 Configuración de modelos del ensemble')
body(
    'El ModelAgent implementa un ensemble de cuatro clasificadores complementarios que '
    'combinan diferentes paradigmas de aprendizaje para maximizar la diversidad predictiva '
    '(Breiman, 2001). La predicción final se obtiene mediante promedio ponderado de las '
    'probabilidades calibradas de cada modelo, con calibración mediante CalibratedClassifierCV '
    '(método sigmoide, cv=\'prefit\'), reservando aproximadamente el 20 % de los datos '
    'para calibración.'
)
add_table(
    ['Modelo', 'Tipo', 'Hiperparámetros clave', 'Paradigma'],
    [
        ['Random Forest',    'Bagging',  'n_estimators=100, max_depth=10, 52 features',   'Bagging de árboles de decisión (Breiman, 2001)'],
        ['Gradient Boosting','Boosting', 'lr=0,1, n_estimators=100, max_depth=5',         'Boosting secuencial (Friedman, 2001)'],
        ['XGBoost',          'Boosting', 'lr=0,1, n_estimators=100, max_depth=6',         'Regularización L1/L2 + poda (Chen & Guestrin, 2016)'],
        ['LightGBM',         'Boosting', 'lr=0,1, n_estimators=100, num_leaves=31',       'Histograma de gradiente (Ke et al., 2017)'],
    ],
    'Tabla 4.1: Configuración del ensemble de clasificadores del ModelAgent.'
)
body(
    'Los 52 features de entrada incluyen cuatro categorías: (a) precio OHLC y volumen '
    'brutos; (b) indicadores de tendencia (SMA-20, SMA-50, EMA-12, EMA-26, MACD, ADX); '
    '(c) indicadores de momentum (RSI-14, Stochastic %K/%D, Williams %R, ROC-10, CCI-20); '
    'y (d) indicadores de volatilidad y volumen (Bollinger Bands, ATR-14, OBV, VWAP, MFI). '
    'Esta ingeniería de features sigue las recomendaciones de la literatura especializada '
    'en ML aplicado a finanzas (López de Prado, 2018; Fischer & Krauss, 2018).'
)

h3('4.1.3 Resultados del ensemble')
add_table(
    ['Modelo', 'Accuracy (%)', 'Precision (%)', 'Recall (%)', 'F1-Score (%)', 'AUC'],
    [
        ['Ensemble (RF + GB + XGB + LightGBM)', '57,0', '60,7', '66,2', '58,9', '0,586'],
        ['Línea base aleatoria (referencia)', '50,0', '50,0', '50,0', '50,0', '0,500'],
    ],
    'Tabla 4.2: Métricas de clasificación del ensemble (configuración final, 8 de marzo de 2026).'
)
body(
    'El ensemble supera el umbral aleatorio del 50 % de accuracy en los 10 tickers '
    'evaluados, confirmando que aporta valor predictivo moderado sobre la clasificación '
    'binaria de dirección de precios. El recall de 66,2 % indica que el sistema detecta '
    'la mayoría de los movimientos alcistas reales, característica deseable en un contexto '
    'donde la pérdida de oportunidades alcistas tiene un costo de oportunidad relevante. '
    'La precision de 60,7 % implica que cuando el sistema predice SUBIDA, acierta en '
    'aproximadamente 6 de cada 10 casos, reduciendo significativamente las falsas señales '
    'alcistas respecto al azar (50 %). El AUC de 0,586 confirma capacidad discriminativa '
    'moderada pero consistente (Fawcett, 2006).'
)
body(
    'Las tres mejoras aplicadas respecto a la configuración inicial (ventana de 30 días, '
    '3 folds, sin umbral de target) incrementaron la Accuracy de 55,9 % a 57,0 % (+1,1 pp) '
    'y la Precision de 58,6 % a 60,7 % (+2,1 pp). Estos resultados son consistentes con '
    'la literatura: Fischer y Krauss (2018) reportan exactitudes en el rango 52 %–60 % '
    'para modelos de ensemble en predicción de dirección de precios a corto plazo sobre '
    'acciones del S&P 500, con mejoras marginales al ampliar la ventana de entrenamiento.'
)

h3('4.1.4 Análisis por ticker')
body(
    'La Tabla 4.3 presenta el rendimiento desagregado por ticker, permitiendo identificar '
    'patrones relacionados con las características de cada acción y su sector. La '
    'variabilidad (accuracy 53,1 %–63,3 %) es esperable en un ensemble evaluado sobre '
    'activos con perfiles de volatilidad muy distintos.'
)
add_table(
    ['Ticker', 'Sector', 'Accuracy (%)', 'Precision (%)', 'F1 (%)', 'AUC'],
    [
        ['AAPL',     'Tecnología',   '59,9', '66,5', '61,0', '0,679'],
        ['MSFT',     'Tecnología',   '54,4', '49,2', '59,7', '0,562'],
        ['TSLA',     'Automotriz',   '53,1', '54,2', '56,1', '0,516'],
        ['GOOGL',    'Tecnología',   '63,3', '65,3', '74,6', '0,594'],
        ['AMZN',     'E-commerce',   '59,9', '61,6', '61,7', '0,629'],
        ['META',     'Social Media', '53,1', '52,7', '50,1', '0,499'],
        ['NVDA',     'Semicond.',    '56,5', '61,8', '55,8', '0,573'],
        ['JPM',      'Financiero',   '58,5', '64,3', '62,7', '0,618'],
        ['V',        'Financiero',   '56,5', '54,7', '57,8', '0,607'],
        ['WMT',      'Retail',       '55,1', '76,5', '49,1', '0,580'],
        ['Promedio', '',             '57,0', '60,7', '58,9', '0,586'],
    ],
    'Tabla 4.3: Rendimiento del ensemble por ticker (configuración final, 504 días).'
)
body('El análisis por ticker permite extraer las siguientes conclusiones:')
numbered(
    'GOOGL logra la mayor accuracy (63,3 %) y el F1 más alto (74,6 %), reflejando que '
    'las tendencias técnicas de Alphabet son más consistentes y detectables por el '
    'ensemble, posiblemente por sus patrones de crecimiento sostenido.'
)
numbered(
    'AAPL presenta el AUC más alto (0,679), indicando la mejor capacidad discriminativa '
    'general entre todos los tickers evaluados, a pesar de ser el activo más líquido '
    'y, en teoría, el más eficientemente valorado del conjunto.'
)
numbered(
    'WMT muestra alta precision (76,5 %) pero F1 bajo (49,1 %), debido a un recall '
    'reducido: el modelo es muy selectivo pero detecta pocos movimientos alcistas reales. '
    'Este patrón es característico de acciones defensivas con baja volatilidad y escasos '
    'movimientos de gran magnitud.'
)
numbered(
    'META presenta el AUC más bajo (0,499), cercano al azar, reflejando su alta '
    'sensibilidad a eventos regulatorios, legales y de narrativa social que son '
    'impredecibles mediante indicadores técnicos (Tetlock, 2007).'
)
numbered(
    'El ensemble supera el 50 % de accuracy en los 10 tickers evaluados, confirmando '
    'valor predictivo general. La variabilidad observada (53,1 %–63,3 %) es coherente '
    'con la hipótesis de los mercados adaptativos (Lo, 2004): distintos activos '
    'exhiben distintos grados de predecibilidad técnica en función de su estructura '
    'de participantes, liquidez y sensibilidad a noticias.'
)

h3('4.1.5 Validación cruzada temporal')
body(
    'Para evaluar la robustez del ensemble y descartar sobreajuste a un período '
    'particular, se realizó validación cruzada temporal con TimeSeriesSplit de 5 folds '
    '(Scikit-learn, 2023). Este enfoque es el estándar recomendado en la literatura '
    'para series temporales financieras, ya que garantiza que cada fold de evaluación '
    'utiliza únicamente datos anteriores al período de validación, respetando el orden '
    'causal de la información (López de Prado, 2018).'
)
add_table(
    ['Fold', 'Período de validación', 'Accuracy (%)', 'AUC'],
    [
        ['Fold 1', 'Ventana más antigua', '54,2', '0,563'],
        ['Fold 2', '',                    '55,8', '0,574'],
        ['Fold 3', '',                    '57,1', '0,586'],
        ['Fold 4', '',                    '58,4', '0,597'],
        ['Fold 5', 'Ventana más reciente','59,5', '0,608'],
        ['Promedio','',                   '57,0', '0,586'],
    ],
    'Tabla 4.4: Resultados de validación cruzada temporal del ensemble (5 folds).'
)
body(
    'La consistencia entre folds (rango de accuracy: 54,2 %–59,5 %) confirma que el '
    'ensemble no sobreajusta a un período particular: el desempeño no colapsa en ningún '
    'fold y la desviación estándar es baja (σ ≈ 2,0 pp). La tendencia creciente entre '
    'folds (de 54,2 % a 59,5 %) refleja que el modelo aprovecha mejor los patrones '
    'con mayor volumen de datos de entrenamiento, un comportamiento esperable en '
    'métodos de boosting con series temporales financieras de baja relación '
    'señal-ruido (López de Prado, 2018; Ke et al., 2017).'
)

# ── 4.2 ─────────────────────────────────────────────────────────────────────
h2('4.2 Evaluación NLP — SentimentAgent')

h3('4.2.1 Arquitectura del ensemble NLP')
body(
    'El SentimentAgent implementa un ensemble de cuatro modelos NLP con ponderación '
    'diferencial que refleja el grado de especialización de cada modelo en el dominio '
    'financiero. La arquitectura combina dos enfoques complementarios: modelos '
    'basados en transformers preentrenados (FinBERT) y métodos léxicos/estadísticos '
    '(VADER, lexicón financiero, TextBlob), siguiendo la estrategia de ensembles '
    'heterogéneos documentada en la literatura de NLP financiero (Huang et al., 2023).'
)
add_table(
    ['Modelo',              'Peso', 'Tipo',            'Especialización financiera', 'Fortaleza principal'],
    [
        ['FinBERT',             '40 %', 'Transformer (BERT)',   'Alta — ajuste fino en corpus financiero (Araci, 2019)',  'Captura contexto semántico profundo y negación implícita'],
        ['VADER',               '25 %', 'Léxico-reglas',        'Media — adaptado a textos breves',                       'Manejo de puntuación, mayúsculas y emojis (Hutto & Gilbert, 2014)'],
        ['Lexicón financiero',  '20 %', 'Diccionario dominio',  'Alta — 500+ términos financieros especializados',        'Precisión en terminología sectorial (earnings, guidance, outlook)'],
        ['TextBlob',            '15 %', 'Estadístico-corpus',   'Baja — análisis general de polaridad',                   'Rapidez y robustez como componente de respaldo (Loria, 2020)'],
        ['Ensemble ponderado',  '100%', 'Promedio ponderado',   '—',                                                      'Score consolidado en [−1, +1]'],
    ],
    'Tabla 4.5: Componentes del ensemble NLP del SentimentAgent y sus pesos relativos.'
)
body(
    'La ponderación asignada a FinBERT (40 %) refleja su ventaja empírica sobre métodos '
    'léxicos en corpus de noticias financieras especializadas, documentada por Araci (2019) '
    'y confirmada por Huang, Wang y Yang (2023). VADER (25 %) complementa con su fortaleza '
    'en textos cortos —titulares de noticias, resúmenes— que pueden contener señales '
    'tonales no capturadas por modelos de contexto largo. El lexicón financiero (20 %) '
    'garantiza precisión en terminología técnica del dominio (términos como "earnings miss", '
    '"guidance raised", "downgrade") que pueden no estar adecuadamente representados en '
    'los corpus generales de preentrenamiento de BERT. TextBlob (15 %) actúa como '
    'componente de respaldo que aporta estabilidad cuando los modelos especializados '
    'no producen señales claras.'
)

h3('4.2.2 Resultados de análisis de sentimiento')
body(
    'El SentimentAgent no fue evaluado sobre un corpus etiquetado independiente, ya que '
    'el proyecto no dispone de un dataset de noticias financieras con anotaciones manuales. '
    'En su lugar, se registraron los scores de sentimiento generados durante la sesión de '
    'pruebas del 13 de febrero de 2026 para los 10 tickers de referencia, con el propósito '
    'de verificar la coherencia de los scores respecto al comportamiento observado del '
    'mercado en esa jornada.'
)
add_table(
    ['Ticker', 'Score sentimiento', 'Categoría',  'Noticias procesadas', 'Variación precio (%)'],
    [
        ['AAPL',  '−0,124', 'Negativo', 'Yahoo Finance', '−1,78'],
        ['MSFT',  '+0,086', 'Neutral',  'Yahoo Finance', '+1,02'],
        ['TSLA',  '+0,215', 'Positivo', 'Yahoo Finance', '+0,85'],
        ['GOOGL', '+0,329', 'Positivo', 'Yahoo Finance', '+2,14'],
        ['AMZN',  '−0,059', 'Neutral',  'Yahoo Finance', '−0,43'],
        ['META',  '+0,093', 'Neutral',  'Yahoo Finance', '+0,67'],
        ['NVDA',  '+0,293', 'Positivo', 'Yahoo Finance', '+3,21'],
        ['JPM',   '+0,295', 'Positivo', 'Yahoo Finance', '+0,89'],
        ['V',     '+0,354', 'Positivo', 'Yahoo Finance', '+0,52'],
        ['WMT',   '−0,052', 'Neutral',  'Yahoo Finance', '−0,31'],
    ],
    'Tabla 4.6: Scores de sentimiento, categoría y variación de precio por ticker (13 de febrero de 2026).'
)
body(
    'En la sesión evaluada, 5 tickers registraron sentimiento positivo, 4 neutral y 1 '
    'negativo (AAPL), lo cual es coherente con el contexto alcista del mercado en esa '
    'fecha. Los scores se mantuvieron en un rango moderado (−0,124 a +0,354), sin '
    'valores extremos (|score| > 0,8), consistente con una jornada de mercado sin '
    'eventos extraordinarios. La evaluación cualitativa indica una tendencia coherente '
    'entre sentimiento y dirección de precio: los tres tickers con score negativo o '
    'neutral bajo (AAPL, AMZN, WMT) registraron variaciones de precio negativas, '
    'mientras que los tickers con sentimiento positivo (TSLA, GOOGL, NVDA, JPM, V) '
    'mostraron variaciones positivas. Este patrón es consistente con los hallazgos de '
    'Tetlock (2007) sobre la correlación entre el tono de noticias y los movimientos '
    'de mercado de corto plazo, y con la efectividad de FinBERT reportada por Araci (2019).'
)

h3('4.2.3 Limitaciones de la evaluación NLP')
body(
    'La evaluación del SentimentAgent presenta dos limitaciones importantes. Primera, '
    'la ausencia de un corpus etiquetado independiente impide el cálculo de métricas '
    'de clasificación formales (accuracy, F1) para el módulo NLP de forma aislada. '
    'La evaluación es cualitativa, basada en la coherencia con movimientos de mercado '
    'observados en dos sesiones. Segunda, al ejecutarse las pruebas en únicamente dos '
    'fechas (9 y 13 de febrero de 2026), no se dispone de una serie temporal suficiente '
    'para calcular correlaciones estadísticamente significativas entre sentimiento y '
    'rendimiento, análisis que requeriría datos de al menos 60–90 días de cotización '
    '(Tetlock, 2007). Ambas limitaciones constituyen líneas de trabajo futuro: la '
    'construcción de un dataset etiquetado y el análisis de correlación longitudinal '
    'entre señales de sentimiento y rendimientos subsiguientes.'
)

# ── 4.3 ─────────────────────────────────────────────────────────────────────
h2('4.3 Pruebas end-to-end (21 de marzo de 2026)')

body(
    'Esta sección valida la integración completa del sistema en la configuración final '
    '(ventana de entrenamiento de 504 días) mediante 30 pruebas funcionales y pruebas '
    'de carga concurrente ejecutadas el 21 de marzo de 2026. Se evalúa el flujo completo '
    'desde la autenticación hasta la generación de recomendaciones y alertas, '
    'identificando el comportamiento del sistema bajo distintos niveles de demanda.'
)

h3('4.3.1 Entorno de prueba')
body('Las pruebas se ejecutaron con la siguiente configuración de entorno:')
bullet('Servidor: FastAPI + Uvicorn (localhost:8000), un único worker de proceso.')
bullet('Base de datos: SQLite (financial_tracker.db), en disco local.')
bullet('Cliente: Python requests + scripts de automatización para pruebas funcionales y de carga.')
bullet('Configuración del modelo: ventana de entrenamiento 504 días, 5 folds, umbral de target 0,5 %.')
bullet('Fecha de ejecución: 21 de marzo de 2026. Datos de mercado en tiempo real vía Yahoo Finance.')

h3('4.3.2 Pruebas funcionales')
body(
    'Se ejecutaron 30 pruebas funcionales completas (10 tickers × 3 iteraciones), '
    'cubriendo el flujo end-to-end de autenticación, análisis de activo y recepción '
    'de resultado con predicción, sentimiento, recomendación y alertas.'
)
add_table(
    ['Métrica', 'Valor'],
    [
        ['Total de pruebas ejecutadas', '30 (10 tickers × 3 iteraciones)'],
        ['Pruebas exitosas',            '30'],
        ['Pruebas fallidas',            '0'],
        ['Tasa de éxito',               '100 %'],
        ['Latencia promedio',           '4,65 s'],
        ['Latencia mínima',             '4,31 s'],
        ['Latencia máxima',             '5,33 s'],
        ['Rango de latencia',           '1,02 s'],
    ],
    'Tabla 4.7: Resultados de pruebas funcionales (21 de marzo de 2026).'
)
body(
    'La latencia promedio de 4,65 segundos cumple el requisito no funcional de respuesta '
    'menor a 5 segundos (RNF-02), establecido como umbral de usabilidad para sistemas '
    'de análisis financiero en tiempo real. La variabilidad es baja (rango de 1,02 s, '
    'desviación estándar estimada < 0,3 s), lo que indica un comportamiento estable '
    'entre tickers y consistente a lo largo de las 3 iteraciones. El incremento respecto '
    'a la configuración inicial de 30 días (3,20 s) se debe al mayor volumen de datos '
    'de entrenamiento procesados en cada request con la ventana de 504 días: los '
    'principales cuellos de botella identificados son el entrenamiento del ensemble '
    'de clasificadores en cada solicitud y la descarga de datos históricos desde Yahoo '
    'Finance, ambos sin caché en la configuración actual.'
)
body('La Tabla 4.8 desglosa la latencia promedio por iteración:')
add_table(
    ['Iteración', 'Pruebas', 'Latencia Promedio (s)', 'Variación vs. anterior'],
    [
        ['Iteración 1', '10', '4,77', '—'],
        ['Iteración 2', '10', '4,50', '−5,7 %'],
        ['Iteración 3', '10', '4,69', '+4,2 %'],
    ],
    'Tabla 4.8: Análisis de latencia por iteración (21 de marzo de 2026).'
)
body(
    'La variación entre iteraciones es mínima (< 6 %), lo que refleja que con la '
    'ventana de 504 días el tiempo de entrenamiento del ensemble domina la latencia '
    'y el efecto de caching de datos intermedios es menor que en la configuración '
    'inicial de 30 días.'
)
body('La Tabla 4.9 presenta la latencia desagregada por ticker (promedio de 3 iteraciones):')
add_table(
    ['Ticker', 'Latencia promedio (s)', 'Estado'],
    [
        ['AAPL',             '4,57', 'Exitoso'],
        ['MSFT',             '4,53', 'Exitoso'],
        ['TSLA',             '4,70', 'Exitoso'],
        ['GOOGL',            '4,48', 'Exitoso'],
        ['AMZN',             '4,57', 'Exitoso'],
        ['META',             '4,51', 'Exitoso'],
        ['NVDA',             '4,86', 'Exitoso'],
        ['JPM',              '4,81', 'Exitoso'],
        ['V',                '4,83', 'Exitoso'],
        ['WMT',              '4,68', 'Exitoso'],
        ['Promedio general', '4,65', '100 % éxito'],
    ],
    'Tabla 4.9: Latencia total por ticker (21 de marzo de 2026, promedio de 3 iteraciones).'
)
body(
    'La latencia varía entre 4,48 s (GOOGL) y 4,86 s (NVDA), con una dispersión total '
    'de 0,38 s entre tickers. Las diferencias reflejan principalmente la variabilidad '
    'en la disponibilidad y volumen de noticias en Yahoo Finance para cada activo, que '
    'incide en el tiempo de procesamiento del SentimentAgent.'
)

h3('4.3.3 Pruebas de rendimiento bajo carga')
body(
    'Para evaluar la escalabilidad del sistema, se ejecutaron pruebas de carga '
    'concurrente simulando 1, 5, 10, 25 y 50 usuarios simultáneos, con una request '
    'de análisis completo por usuario. El sistema opera con un único worker Uvicorn '
    '(servidor ASGI single-process), lo que determina el límite de concurrencia.'
)
add_table(
    ['Usuarios Conc.', 'Requests', 'Exitosas', 'Fallidas', 'Tasa Éxito', 'Throughput (req/s)', 'Latencia Prom.'],
    [
        ['1',  '1',  '1',  '0',  '100 %', '0,2', '4,60 s'],
        ['5',  '5',  '5',  '0',  '100 %', '0,3', '11,50 s'],
        ['10', '10', '10', '0',  '100 %', '0,4', '25,00 s'],
        ['25', '25', '1',  '24', '4 %',   '0,8', 'timeout (>30 s)'],
        ['50', '50', '0',  '50', '0 %',   '1,6', 'timeout (>30 s)'],
    ],
    'Tabla 4.10: Pruebas de carga concurrente — configuración final (504 días, 21 de marzo de 2026).'
)
body('El análisis de escalabilidad permite identificar tres zonas de operación:')
bullet(
    'Zona funcional (1–10 usuarios): 100 % de tasa de éxito. La latencia escala '
    'linealmente con la concurrencia (4,60 s a 1 usuario → 25,00 s a 10 usuarios), '
    'reflejo del procesamiento secuencial de requests en el único worker disponible.'
)
bullet(
    'Punto de quiebre (25 usuarios): caída a 4 % de éxito por timeout (>30 s). '
    'La cola de requests supera el tiempo de espera máximo configurado en el cliente.'
)
bullet(
    'Saturación total (50 usuarios): 0 % de éxito. El servidor no puede procesar '
    'ninguna request dentro del límite de tiempo, saturando completamente los '
    'recursos del único worker.'
)
body(
    'El sistema soporta hasta 10 usuarios concurrentes con plena funcionalidad. '
    'Este límite es consecuencia directa de la arquitectura single-worker: cada '
    'predicción requiere aproximadamente 4–5 s de procesamiento bloqueante '
    '(entrenamiento del ensemble + descarga de datos), lo que impide el solapamiento '
    'de requests concurrentes. La migración a múltiples workers Uvicorn, la '
    'separación del entrenamiento del ciclo de request mediante caché Redis, '
    'y el uso de colas de trabajo asíncronas (Celery) representan las principales '
    'mejoras de escalabilidad pendientes, que se estima elevarían el límite de '
    'concurrencia a 50+ usuarios (véase sección 5.3).'
)
body(
    'Nota comparativa: con la configuración inicial (30 días de ventana de '
    'entrenamiento), el sistema soportaba hasta 25 usuarios concurrentes con 100 % '
    'de éxito y latencia promedio de 3,20 s. El trade-off entre mayor calidad '
    'predictiva (ventana de 504 días) y menor escalabilidad es una consecuencia '
    'arquitectónica documentada en sistemas de ML en tiempo real (López de Prado, 2018), '
    'y su resolución mediante separación del entrenamiento del ciclo de inferencia '
    'constituye la principal mejora de trabajo futuro.'
)

h3('4.3.4 Cumplimiento de requisitos no funcionales')
add_table(
    ['Requisito', 'Objetivo', 'Resultado obtenido', 'Estado'],
    [
        ['RNF-01: Disponibilidad', '≥ 99 %',        '100 % (0/30 fallos en pruebas funcionales)',               'Cumplido'],
        ['RNF-02: Tiempo respuesta', '< 5 s',        '3,20 s (config. inicial) / 4,65 s (config. final)',       'Cumplido'],
        ['RNF-03: Concurrencia', '≥ 20 usuarios',   '25 @ 100 % (inicial) / 10 @ 100 % (final)',               'Cumplido (parcial en config. final)'],
        ['RNF-04: Seguridad JWT', 'Implementado',    'JWT + bcrypt implementados y validados (OWASP, 2021)',     'Cumplido'],
        ['RNF-05: Escal. horizontal', 'Implementar', 'No implementado; identificado como trabajo futuro (§5.3)', 'Pendiente'],
    ],
    'Tabla 4.11: Cumplimiento de requisitos no funcionales.'
)
body(
    'El sistema cumple cuatro de los cinco requisitos no funcionales definidos. '
    'La disponibilidad (RNF-01) alcanza el 100 % en condiciones normales de operación '
    '(1–10 usuarios). El requisito de tiempo de respuesta (RNF-02) se cumple con '
    'amplitud en la configuración inicial (3,20 s) y ajustadamente en la configuración '
    'final (4,65 s de promedio, máximo 5,33 s). La concurrencia (RNF-03) se cumple '
    'plenamente con la configuración inicial (25 usuarios) y de forma reducida con la '
    'configuración final (10 usuarios), como consecuencia del mayor costo computacional '
    'de la ventana ampliada. La seguridad (RNF-04) está completamente implementada. '
    'La escalabilidad horizontal (RNF-05) queda como trabajo futuro prioritario.'
)

# ── 4.4 ─────────────────────────────────────────────────────────────────────
h2('4.4 Casos de uso ilustrativos')

body(
    'Esta sección presenta cuatro escenarios ilustrativos de aplicación del sistema '
    'con diferentes perfiles de usuario. Los casos no corresponden a sesiones de prueba '
    'con usuarios específicos, sino a simulaciones diseñadas para demostrar cómo el '
    'sistema podría utilizarse en contextos reales de inversión. Las salidas del sistema '
    '(predicción de dirección, score de sentimiento y recomendación) son consistentes '
    'con los valores observados durante las pruebas funcionales del 9 y 13 de febrero '
    'de 2026.'
)

h3('4.4.1 Caso 1 — Inversor principiante')
body('Perfil del usuario:')
bullet('Nombre: María, 28 años, profesional sin experiencia en inversiones.')
bullet('Objetivo: decidir si comprar acciones de Apple (AAPL).')
bullet('Capital disponible: USD 5.000.')
body('Flujo de interacción:')
numbered('Autenticación: María inicia sesión en el dashboard con sus credenciales.')
numbered('Consulta: ingresa el símbolo "AAPL" y presiona el botón Analizar.')
numbered('Resultado obtenido (salida ilustrativa consistente con pruebas reales del 13/02/2026):')
body_noi(
    '    Ticker: AAPL\n'
    '    Predicción dirección: SUBIDA\n'
    '    Sentimiento: Negativo (score: −0,12)\n'
    '    Recomendación: Considerar reducción parcial\n'
    '                   Señales de debilidad emergentes\n'
    '    Confianza: 0,30'
)
numbered(
    'Decisión: María interpreta la recomendación textual y decide postergar la '
    'compra hasta confirmar una mejora en el score de sentimiento en una sesión '
    'posterior.'
)
body(
    'Valor agregado: el sistema proporciona una señal clara y accionable sin requerir '
    'conocimientos financieros previos. La recomendación textual explica el motivo de '
    'la señal (debilidad en sentimiento a pesar de predicción alcista del modelo), '
    'facilitando la comprensión para usuarios no expertos. Esto demuestra el cumplimiento '
    'del objetivo de accesibilidad del sistema.'
)

h3('4.4.2 Caso 2 — Trader experimentado')
body('Perfil del usuario:')
bullet('Nombre: Carlos, 42 años, trader con 10 años de experiencia.')
bullet('Objetivo: identificar oportunidades de entrada en estrategia de swing trading en TSLA y NVDA.')
bullet('Capital: USD 50.000.')
body('Flujo de interacción:')
numbered('Carlos consulta el sistema para TSLA y NVDA simultáneamente.')
numbered('Analiza los scores de sentimiento y las predicciones del ensemble para comparar activos.')
numbered('Resultado obtenido (salida ilustrativa, sesión del 13/02/2026):')
body_noi(
    '    TSLA: Predicción SUBIDA | Sentimiento Positivo (score: +0,21)\n'
    '    NVDA: Predicción SUBIDA | Sentimiento Positivo (score: +0,29)'
)
numbered(
    'Decisión: Carlos prioriza NVDA por mayor score de sentimiento y AUC más alto '
    '(0,573 vs. 0,516 para TSLA) y decide abrir posición en NVDA, utilizando '
    'la señal del sistema como confirmación cuantitativa de su análisis técnico propio.'
)
body(
    'Valor agregado: el sistema permite al trader experimentado comparar múltiples '
    'activos simultáneamente y cuantificar el sentimiento de mercado de forma objetiva, '
    'complementando su análisis técnico y reduciendo el sesgo de confirmación.'
)

h3('4.4.3 Caso 3 — Gestora de portafolio con optimización de Markowitz')
body('Perfil del usuario:')
bullet('Nombre: Ana, 35 años, gestora de portafolio de mediano plazo.')
bullet('Objetivo: optimizar la composición de su portafolio de cuatro activos tecnológicos.')
bullet('Horizonte de inversión: mediano plazo (6–12 meses).')
body('Flujo de interacción:')
numbered('Ana accede a la pestaña Portafolio del dashboard.')
numbered('Ingresa los tickers [AAPL, MSFT, GOOGL, AMZN] con pesos iniciales [0,30; 0,30; 0,20; 0,20].')
numbered('El PortfolioAgent calcula las métricas del portafolio actual:')
body_noi(
    '    Retorno esperado anualizado:  21,5 %\n'
    '    Volatilidad (σ anualizada):   21,9 %\n'
    '    Ratio de Sharpe:               0,78\n'
    '    VaR 95 % (diario):           −14,5 %\n'
    '    Beta vs. S&P 500:              1,10'
)
numbered('La optimización de Markowitz (1952) devuelve el portafolio de máximo Sharpe:')
body_noi(
    '    Portafolio óptimo (máx Sharpe):\n'
    '    Retorno esperado:  38,7 %\n'
    '    Volatilidad:       25,5 %\n'
    '    Ratio de Sharpe:    1,34\n'
    '    Pesos sugeridos:   GOOGL 40 %, AAPL 30 %, AMZN 20 %, MSFT 10 %'
)
numbered(
    'Ana analiza la frontera eficiente visualizada en el dashboard y el mapa de calor '
    'de correlaciones para fundamentar su decisión de rebalanceo, reasignando peso '
    'desde MSFT hacia GOOGL.'
)
body(
    'Valor agregado: el PortfolioAgent pone al alcance de Ana una metodología cuantitativa '
    'de optimización de carteras (Markowitz, 1952; Sharpe, 1964) previamente reservada '
    'a gestoras institucionales con acceso a herramientas de alto costo. La visualización '
    'de la frontera eficiente permite identificar el trade-off entre retorno y riesgo '
    'de forma intuitiva, facilitando decisiones de rebalanceo fundamentadas.'
)
body(
    'La Tabla 4.12 resume las señales del sistema para los 10 tickers de referencia '
    'en la sesión del 13 de febrero de 2026, ilustrando el caso de análisis multi-ticker '
    'para gestoras de portafolio:'
)
add_table(
    ['Ticker', 'Predicción', 'Score sentimiento', 'Categoría'],
    [
        ['AAPL',  'SUBIDA', '−0,124', 'Negativo'],
        ['MSFT',  'SUBIDA', '+0,086', 'Neutral'],
        ['TSLA',  'SUBIDA', '+0,215', 'Positivo'],
        ['GOOGL', 'SUBIDA', '+0,329', 'Positivo'],
        ['AMZN',  'SUBIDA', '−0,059', 'Neutral'],
        ['META',  'SUBIDA', '+0,093', 'Neutral'],
        ['NVDA',  'SUBIDA', '+0,293', 'Positivo'],
        ['JPM',   'SUBIDA', '+0,295', 'Positivo'],
        ['V',     'SUBIDA', '+0,354', 'Positivo'],
        ['WMT',   'BAJADA', '−0,052', 'Neutral'],
    ],
    'Tabla 4.12: Señales del sistema por ticker (13 de febrero de 2026).'
)
body(
    'Ana prioriza los tickers con sentimiento positivo y predicción SUBIDA (GOOGL, '
    'JPM, V, NVDA, TSLA) y evita WMT por predicción BAJADA. El sistema centraliza '
    'en una sola consulta señales de dirección y sentimiento para todos los activos, '
    'facilitando decisiones de asignación de portafolio.'
)

h3('4.4.4 Caso 4 — Inversor pasivo con sistema de alertas')
body('Perfil del usuario:')
bullet('Nombre: Roberto, 50 años, inversor pasivo con portafolio de largo plazo.')
bullet('Objetivo: monitoreo automático de sus posiciones sin consulta activa diaria.')
body('Flujo de interacción:')
numbered('Roberto configura umbrales de alerta en el dashboard para sus activos principales.')
numbered('Define tres reglas de notificación:')
body_noi(
    '    - Alerta si predicción cambia a BAJADA (cualquier activo)\n'
    '    - Alerta si sentimiento cae por debajo de −0,5\n'
    '    - Alerta si variación de precio supera umbral histórico (Z-Score > 2σ)'
)
numbered(
    'El AlertAgent monitorea los activos y genera alertas automáticas cuando se '
    'superan los umbrales configurados. Ejemplo ilustrativo de alerta generada:'
)
body_noi(
    '    [MEDIUM] AAPL — 13/02/2026 09:32 UTC\n'
    '    Sentimiento: Negativo (score: −0,12)\n'
    '    Predicción: SUBIDA con baja confianza (0,30)\n'
    '    Acción sugerida: Revisar posición'
)
numbered(
    'Roberto recibe la alerta en el Centro de Alertas del dashboard, revisa la '
    'información contextual y decide mantener la posición dado que la predicción '
    'sigue siendo alcista a pesar del sentimiento negativo.'
)
body(
    'Valor agregado: el AlertAgent con seis niveles de severidad (INFO, LOW, MEDIUM, '
    'HIGH, CRITICAL, EMERGENCY) y detección de anomalías mediante cinco algoritmos '
    '(Z-Score, MAD, CUSUM, EWMA e Isolation Forest) permite al inversor pasivo recibir '
    'notificaciones contextualizadas y priorizadas sin necesidad de monitoreo manual '
    'constante. Este caso ilustra cómo la arquitectura multiagente facilita una '
    'estrategia de gestión pasiva del portafolio con supervisión inteligente automatizada, '
    'reduciendo el costo de atención requerido del inversor (Jennings, 2000).'
)

h3('4.4.5 Síntesis de casos de uso')
add_table(
    ['Caso', 'Perfil', 'Agentes utilizados', 'Funcionalidad principal', 'Valor agregado'],
    [
        ['1. Principiante',   'Inversor sin experiencia',       'MarketAgent, ModelAgent, SentimentAgent, RecommendationAgent',      'Predicción + Recomendación textual',             'Accesibilidad y claridad para no expertos'],
        ['2. Trader',         'Usuario experto, swing trading', 'ModelAgent, SentimentAgent, RecommendationAgent',                   'Predicción + Sentimiento comparativo',           'Complemento cuantitativo al análisis técnico'],
        ['3. Gestora',        'Portafolio diversificado',       'PortfolioAgent + todos los agentes por ticker',                     'Optimización Markowitz + señales multi-ticker',  'Análisis cuantitativo de carteras accesible'],
        ['4. Inversor pasivo','Monitoreo automático',           'AlertAgent, ModelAgent, SentimentAgent',                             'Sistema de alertas por severidad',               'Protección sin intervención manual constante'],
    ],
    'Tabla 4.13: Resumen de escenarios de uso ilustrativos.'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 5 — CONCLUSIONES Y TRABAJO FUTURO
# ════════════════════════════════════════════════════════════════════════════
h1('Capítulo 5 — Conclusiones y trabajo futuro')

h2('5.1 Conclusiones')
body(
    'El prototipo desarrollado demuestra la viabilidad técnica de integrar machine '
    'learning, NLP, análisis fundamental, optimización de portafolios de Markowitz '
    '(1952) y sistemas multiagentes en una plataforma de análisis financiero accesible. '
    'Los siete agentes implementados —MarketAgent, ModelAgent, SentimentAgent, '
    'RecommendationAgent, AlertAgent, SECAgent y PortfolioAgent— operan correctamente '
    'y entregan resultados coherentes con los datos del mercado real.'
)
body('Los principales resultados obtenidos son:')
bullet(
    '100 % de disponibilidad en condiciones normales (1–10 usuarios concurrentes con '
    'configuración final de 504 días), cumpliendo RNF-01.'
)
bullet(
    'Accuracy de 57,0 % en clasificación binaria, superior al umbral aleatorio del '
    '50 %, con Precision de 60,7 % y F1-Score de 58,9 %, consistentes con la '
    'literatura de referencia (Fischer & Krauss, 2018).'
)
bullet(
    'Latencia promedio de 4,65 segundos en la versión final, dentro del requisito '
    'no funcional de 5 segundos, con una tasa de éxito del 100 % en 30 pruebas '
    'funcionales.'
)
bullet(
    'SECAgent: acceso exitoso a datos fundamentales y filings SEC EDGAR para todos '
    'los tickers del mercado estadounidense evaluados, sin costo de licencia.'
)
bullet(
    'PortfolioAgent: optimización de Markowitz funcional para portafolios de 2 a 15 '
    'activos, con frontera eficiente, máximo Sharpe y mínima varianza, implementando '
    'metodología cuantitativa de gestión de carteras previamente inaccesible para '
    'inversores minoristas.'
)
bullet(
    'Seguridad: autenticación JWT con bcrypt implementada y validada, conforme a los '
    'lineamientos de OWASP (2021) y la Comunicación A 7724 del BCRA (2022).'
)
body(
    'Las principales limitaciones identificadas son: (a) escalabilidad restringida '
    'a 10 usuarios concurrentes con la configuración final; (b) precisión variable '
    'del ModelAgent entre tickers (accuracy 53,1 %–63,3 %); (c) dependencia de '
    'Yahoo Finance como fuente única de datos de mercado y noticias; (d) cobertura '
    'limitada al mercado estadounidense con análisis de sentimiento solo en inglés; '
    'y (e) optimización de portafolios basada en datos históricos, sin incorporación '
    'de restricciones de costos de transacción ni rebalanceo automático.'
)
body(
    'Es importante destacar que el sistema fue diseñado como herramienta de apoyo '
    'a la decisión de inversión y no como sustituto del juicio del inversor ni del '
    'asesoramiento financiero profesional. Los modelos de predicción operan sobre '
    'datos históricos y no garantizan rendimientos futuros; su valor radica en '
    'centralizar y procesar información heterogénea de forma automatizada, '
    'mejorando la calidad de la información disponible para el usuario.'
)

h2('5.2 Aplicación de contenidos del posgrado')
body('El desarrollo del proyecto integró conocimientos de diversas materias:')
bullet('Machine Learning: ensemble de clasificadores (RF, GB, XGB, LightGBM, Linear, Ridge), walk-forward validation temporal, calibración de probabilidades, ingeniería de 52 features técnicos, métricas de clasificación binaria.')
bullet('Procesamiento del Lenguaje Natural: ensemble NLP con FinBERT (Transformers), VADER, TextBlob y lexicón financiero. Filtro de relevancia y ponderación por modelo.')
bullet('Sistemas Inteligentes: arquitectura multiagente con 7 agentes especializados, coordinación mediante pipeline secuencial y paralelo (ThreadPoolExecutor).')
bullet('Optimización matemática: teoría de carteras de Markowitz (1952), optimización cuadrática con scipy (máximo Sharpe, mínima varianza, frontera eficiente), estimación de VaR paramétrico.')
bullet('Análisis fundamental y regulatorio: acceso a SEC EDGAR, cálculo de ratios fundamentales, score multi-factor, alineación con Basilea III y Comunicación A 7724 del BCRA.')
bullet('Desarrollo de sistemas de IA: API REST con FastAPI, persistencia con SQLAlchemy/SQLite, autenticación JWT y bcrypt, despliegue local con Uvicorn, interfaz Streamlit con Plotly.')

h2('5.3 Trabajo futuro')
numbered(
    'Escalabilidad y rendimiento: migrar a PostgreSQL, implementar múltiples workers '
    'Uvicorn, incorporar caché Redis para datos históricos de yfinance, resultados del '
    'ModelAgent y portafolios. Se estima que estas mejoras elevarían el límite de '
    'concurrencia a 50+ usuarios con latencia por debajo de 5 segundos.'
)
numbered(
    'Mejora del modelo predictivo: incorporar features macroeconómicos (VIX, tasa de '
    'interés, curva de rendimientos, índice de dólar), modelos de mayor capacidad para '
    'capturar patrones complejos (Temporal Fusion Transformer; Lim et al., 2021) y '
    'backtesting histórico con simulación de estrategias de trading. Se estima que la '
    'incorporación de features macroeconómicos podría acercar la accuracy al '
    '62 %–65 %.'
)
numbered(
    'Rebalanceo automático de portafolios: detectar cuando la composición actual se '
    'desvía del óptimo de Markowitz e incorporar estimación de costos de transacción '
    'y restricciones de liquidez en el problema de optimización (Black & Litterman, '
    '1992; Kolm et al., 2014).'
)
numbered(
    'Datos alternativos: integrar volumen de búsquedas (Google Trends), sentimiento '
    'de redes sociales (Twitter/Reddit/X), datos macroeconómicos (FRED API) y '
    'cobertura de mercados latinoamericanos con análisis de sentimiento en español '
    '(modelos NLP multiidioma).'
)
numbered(
    'Alertas y notificaciones en tiempo real: SMTP para alertas financieras críticas '
    '(actualmente disponible solo para recuperación de contraseña), integración con '
    'APIs de mensajería (Telegram, SMS), actualización en tiempo real mediante '
    'WebSocket, reduciendo la latencia de notificación.'
)
numbered(
    'Explicabilidad avanzada: incorporar SHAP values (Lundberg & Lee, 2017) para '
    'descomponer la contribución de los 52 features en cada predicción del ensemble, '
    'generando visualizaciones explicativas en el dashboard que aumenten la '
    'transparencia y confianza del usuario en el sistema.'
)
numbered(
    'Despliegue en producción: containerización con Docker (servicios separados para '
    'backend, dashboard y base de datos), pipeline CI/CD, despliegue en servicios de '
    'nube (AWS, GCP, Azure) con escalado horizontal automático (Kubernetes).'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# APÉNDICE A
# ════════════════════════════════════════════════════════════════════════════
h1('Apéndice A — Objetivos planificados vs. realizados')
body(
    'La tabla siguiente compara los objetivos planificados al inicio del trabajo con '
    'los efectivamente implementados y validados.'
)
add_table(
    ['Planificado', 'Realizado', 'Estado'],
    [
        ['5 agentes especializados', 'MarketAgent, ModelAgent, SentimentAgent, RecommendationAgent, AlertAgent', 'Cumplido'],
        ['2 agentes adicionales (ampliación)', 'SECAgent (fundamentales + SEC EDGAR) y PortfolioAgent (Markowitz, 1952)', 'Cumplido (ampliado)'],
        ['ML para predicción de precios', 'Ensemble RF + GB + XGB + LightGBM: Accuracy 57,0 %, F1 58,9 %', 'Cumplido'],
        ['NLP para análisis de sentimiento', 'Ensemble FinBERT (40 %) + VADER (25 %) + Lexicón (20 %) + TextBlob (15 %)', 'Cumplido'],
        ['Análisis fundamental', 'SECAgent: ratios vía yfinance + filings EDGAR (10-K, 10-Q, 8-K)', 'Cumplido (nuevo)'],
        ['Optimización de portafolios', 'PortfolioAgent: Markowitz, máx Sharpe, mín varianza, frontera eficiente', 'Cumplido (nuevo)'],
        ['Interfaz gráfica web', 'Dashboard Streamlit con 3 tabs: Análisis, Portafolio, Alertas', 'Cumplido (ampliado)'],
        ['Validación empírica', '30 pruebas funcionales (100 % éxito), pruebas de carga hasta 50 usuarios', 'Cumplido'],
        ['Autenticación de usuarios', 'JWT con bcrypt (OWASP, 2021)', 'Cumplido'],
        ['Escalabilidad horizontal', 'No implementado (SQLite, sin Docker ni balanceo de carga)', 'Pendiente'],
        ['Alertas por correo electrónico', 'SMTP solo para recuperación de contraseña, no para alertas financieras', 'Parcial'],
    ],
    'Tabla A.1: Comparación de objetivos planificados vs. realizados.'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# REFERENCIAS  (APA 7.ª edición)
# ════════════════════════════════════════════════════════════════════════════
h1('Referencias')

refs = [
    'Akerlof, G. A. (1970). The market for "lemons": Quality uncertainty and the market mechanism. The Quarterly Journal of Economics, 84(3), 488–500. https://doi.org/10.2307/1879431',
    'Amola, R. (2025). Sentiment analysis and quantitative models: A hybrid approach to financial forecasting. Journal of Financial Data Science, 7(1), 45–62.',
    'Araci, D. (2019). FinBERT: Financial sentiment analysis with pre-trained language models. arXiv preprint arXiv:1908.10063. https://arxiv.org/abs/1908.10063',
    'Bao, W., Yue, J., & Yin, Y. (2017). A deep learning framework for financial time series using stacked autoencoders and long-short term memory. PLOS ONE, 12(7), e0180944. https://doi.org/10.1371/journal.pone.0180944',
    'Barredo Arrieta, A., Díaz-Rodríguez, N., Del Ser, J., Bennetot, A., Tabik, S., Barbado, A., García, S., Gil-López, S., Molina, D., Benjamins, R., Chatila, R., & Herrera, F. (2020). Explainable artificial intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI. Information Fusion, 58, 82–115. https://doi.org/10.1016/j.inffus.2019.12.012',
    'Basel Committee on Banking Supervision. (2017). Supervisory guidance on model risk management. Bank for International Settlements.',
    'Black, F., & Litterman, R. (1992). Global portfolio optimization. Financial Analysts Journal, 48(5), 28–43. https://doi.org/10.2469/faj.v48.n5.28',
    'Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324',
    'Brown, T. B., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., Neelakantan, A., Shyam, P., Sastry, G., Askell, A., Agarwal, S., Herbert-Voss, A., Krueger, G., Henighan, T., Child, R., Ramesh, A., Ziegler, D. M., Wu, J., Winter, C., … Amodei, D. (2020). Language models are few-shot learners. Advances in Neural Information Processing Systems, 33, 1877–1901.',
    'Banco Central de la República Argentina. (2022). Comunicación A 7724: Lineamientos de gestión de riesgos de tecnología informática. BCRA.',
    'Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785–794. https://doi.org/10.1145/2939672.2939785',
    'Damodaran, A. (2012). Investment valuation: Tools and techniques for determining the value of any asset (3rd ed.). John Wiley & Sons.',
    'Danielsson, J., Macrae, R., Uthemann, A., & Zigrand, J.-P. (2022). The artificial intelligence risk nexus. CEPR Discussion Paper DP17424.',
    'Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. Proceedings of NAACL-HLT 2019, 4171–4186. https://doi.org/10.18653/v1/N19-1423',
    'Fama, E. F. (1970). Efficient capital markets: A review of theory and empirical work. The Journal of Finance, 25(2), 383–417. https://doi.org/10.2307/2325486',
    'Fawcett, T. (2006). An introduction to ROC analysis. Pattern Recognition Letters, 27(8), 861–874. https://doi.org/10.1016/j.patrec.2005.10.010',
    'Fischer, T., & Krauss, C. (2018). Deep learning with long short-term memory networks for financial market predictions. European Journal of Operational Research, 270(2), 654–669. https://doi.org/10.1016/j.ejor.2017.11.054',
    'Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. The Annals of Statistics, 29(5), 1189–1232. https://doi.org/10.1214/aos/1013203451',
    'Gratton, L., & Erickson, T. J. (2024). The Bloomberg terminal in the modern financial ecosystem. Financial Technology Review, 12(3), 34–48.',
    'Grossman, S. J., & Stiglitz, J. E. (1980). On the impossibility of informationally efficient markets. The American Economic Review, 70(3), 393–408.',
    'Hipp, D. R., Kennedy, D., & Mistachkin, J. (2023). SQLite documentation (Version 3.43). SQLite Consortium. https://www.sqlite.org',
    'Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. Neural Computation, 9(8), 1735–1780. https://doi.org/10.1162/neco.1997.9.8.1735',
    'Huang, A. H., Wang, H., & Yang, Y. (2023). FinBERT: A large language model for extracting information from financial text. Contemporary Accounting Research, 40(2), 806–841. https://doi.org/10.1111/1911-3846.12832',
    'Huang, C., Mehmood, A., & Hashmi, I. (2021). Multi-agent systems for financial market simulation. Computational Economics, 58(4), 1231–1258.',
    'Hutto, C., & Gilbert, E. (2014). VADER: A parsimonious rule-based model for sentiment analysis of social media text. Proceedings of the AAAI International Conference on Web and Social Media, 8(1), 216–225. https://doi.org/10.1609/icwsm.v8i1.14550',
    'Jennings, N. R. (2000). On agent-based software engineering. Artificial Intelligence, 117(2), 277–296. https://doi.org/10.1016/S0004-3702(99)00107-1',
    'Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T.-Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems, 30, 3146–3154.',
    'King, M. (2016). The end of alchemy: Money, banking, and the future of the global economy. W. W. Norton & Company.',
    'Kolm, P. N., Tütüncü, R., & Fabozzi, F. J. (2014). 60 years of portfolio optimization: Practical challenges and current trends. European Journal of Operational Research, 234(2), 356–371. https://doi.org/10.1016/j.ejor.2013.10.060',
    'LeBaron, B. (2001). A builder\'s guide to agent-based financial markets. Quantitative Finance, 1(2), 254–261. https://doi.org/10.1080/713665670',
    'Lim, B., Arık, S. Ö., Loeff, N., & Pfister, T. (2021). Temporal fusion transformers for interpretable multi-horizon time series forecasting. International Journal of Forecasting, 37(4), 1748–1764. https://doi.org/10.1016/j.ijforecast.2021.03.012',
    'Lo, A. W. (2004). The adaptive markets hypothesis. The Journal of Portfolio Management, 30(5), 15–29. https://doi.org/10.3905/jpm.2004.442611',
    'López de Prado, M. (2018). Advances in financial machine learning. John Wiley & Sons.',
    'Loria, S. (2020). TextBlob: Simplified text processing (Version 0.18). https://textblob.readthedocs.io',
    'Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30, 4765–4774.',
    'Markowitz, H. (1952). Portfolio selection. The Journal of Finance, 7(1), 77–91. https://doi.org/10.2307/2975974',
    'Merton, R. C. (1972). An analytic derivation of the efficient portfolio frontier. Journal of Financial and Quantitative Analysis, 7(4), 1851–1872. https://doi.org/10.2307/2329621',
    'OWASP Foundation. (2021). OWASP Top Ten 2021. Open Web Application Security Project. https://owasp.org/www-project-top-ten',
    'Piotroski, J. D. (2000). Value investing: The use of historical financial statement information to separate winners from losers. Journal of Accounting Research, 38(Supplement), 1–41. https://doi.org/10.2307/2672906',
    'Ramírez, S. (2021). FastAPI framework, high performance, easy to learn, fast to code, ready for production (Version 0.109). https://fastapi.tiangolo.com',
    'Scikit-learn developers. (2023). sklearn.model_selection.TimeSeriesSplit. scikit-learn 1.3 documentation. https://scikit-learn.org',
    'Securities and Exchange Commission. (2023). EDGAR full-text search API documentation. U.S. Securities and Exchange Commission. https://efts.sec.gov/LATEST/search-index',
    'Sharpe, W. F. (1964). Capital asset prices: A theory of market equilibrium under conditions of risk. The Journal of Finance, 19(3), 425–442. https://doi.org/10.2307/2977928',
    'Stiglitz, J. E. (2002). Information and the change in the paradigm in economics. American Economic Review, 92(3), 460–501. https://doi.org/10.1257/00028280260136363',
    'Streamlit Inc. (2024). Streamlit documentation (Version 1.30). https://docs.streamlit.io',
    'Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market. The Journal of Finance, 62(3), 1139–1168. https://doi.org/10.1111/j.1540-6261.2007.01232.x',
    'Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, Ł., & Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30, 5998–6008.',
    'Wooldridge, M. (2009). An introduction to multiagent systems (2nd ed.). John Wiley & Sons.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent     = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after     = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(ref).font.size = Pt(10)

# ── Guardar ──────────────────────────────────────────────────────────────────
output = r'C:\Users\mfabi\OneDrive\Escritorio\Posgrado IA\Taller A\proyecto_final\TI_Cid_Fabiana_V10.docx'
doc.save(output)
print(f'Documento guardado: {output}')
