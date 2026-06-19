"""
Fusiona V15_anterior (estructura completa con índice) con el contenido nuevo de V15_current
(SHAP/MDI/MDA, HRP, BacktestAgent mejorado).
"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

STYLE_MAP = {
    'Normal': 'Normal',
    'Heading 1': 'Ttulo1',
    'Heading 2': 'Ttulo2',
    'Heading 3': 'Ttulo3',
    'List Bullet': 'ListBullet',
    'List Number': 'ListNumber',
    'toc 3': 'TDC3',
}

def replace_text_keep_format(p, new_text):
    """Reemplaza el texto de un párrafo manteniendo el rPr del primer run."""
    rpr_xml = None
    for r in p.runs:
        rpr = r._r.find(qn('w:rPr'))
        if rpr is not None:
            rpr_xml = copy.deepcopy(rpr)
        break
    para_el = p._element
    for r in para_el.findall(qn('w:r')):
        para_el.remove(r)
    for hl in para_el.findall(qn('w:hyperlink')):
        para_el.remove(hl)
    r_el = OxmlElement('w:r')
    if rpr_xml is not None:
        r_el.append(rpr_xml)
    t_el = OxmlElement('w:t')
    t_el.text = new_text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_el.append(t_el)
    para_el.append(r_el)

def insert_after_element(ref_el, text, style='Normal'):
    """Inserta un nuevo párrafo w:p DESPUÉS del elemento XML ref_el en el body."""
    body = ref_el.getparent()
    ref_pos = list(body).index(ref_el)

    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), STYLE_MAP.get(style, style))
    pPr.append(pStyle)
    new_p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    new_p.append(r)
    body.insert(ref_pos + 1, new_p)
    return new_p

def insert_before_element(ref_el, text, style='Normal'):
    """Inserta un nuevo párrafo w:p ANTES del elemento XML ref_el en el body."""
    body = ref_el.getparent()
    ref_pos = list(body).index(ref_el)

    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), STYLE_MAP.get(style, style))
    pPr.append(pStyle)
    new_p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    new_p.append(r)
    body.insert(ref_pos, new_p)
    return new_p

def find_para(doc, search_text, style_hint=None):
    """Busca un párrafo por contenido parcial. Retorna el párrafo o None."""
    for p in doc.paragraphs:
        if search_text in p.text:
            if style_hint is None or style_hint in p.style.name:
                return p
    return None


def main():
    doc = Document('TI_Cid_Fabiana_V15_anterior.docx')
    print("=== INICIANDO FUSIÓN V15 ===")

    # ─── 1. RESUMEN: actualizar keywords ───
    p = find_para(doc, 'Palabras clave: sistema multiagente')
    if p:
        replace_text_keep_format(p,
            'Palabras clave: sistema multiagente, machine learning, SHAP values, MDI, MDA, '
            'análisis de sentimiento, FinBERT, procesamiento del lenguaje natural, '
            'optimización de portafolios, Markowitz, HRP, backtesting walk-forward, '
            'FastAPI, Streamlit, SEC EDGAR.')
        print("  [OK] Palabras clave actualizadas (SHAP, MDI, MDA, HRP)")

    # ─── 2. RESUMEN: actualizar párrafo de integración ───
    p = find_para(doc, 'El trabajo integra conocimientos de aprendizaje autom')
    if p:
        replace_text_keep_format(p,
            'El trabajo integra conocimientos de aprendizaje automático (incluyendo explicabilidad '
            'mediante SHAP values), procesamiento del lenguaje natural, optimización matemática, '
            'análisis fundamental de empresas y seguridad informática, con el objetivo de construir '
            'un sistema de soporte a la decisión accesible, reproducible y validado empíricamente.')
        print("  [OK] Resumen integración actualizado (SHAP values)")

    # ─── 3. SECCIÓN 2.2: actualizar párrafo ML y agregar MDI/MDA y SHAP ───
    p = find_para(doc, 'El aprendizaje automático ha demostrado gran potencial')
    if p:
        replace_text_keep_format(p,
            'El aprendizaje automático ha demostrado gran potencial para modelar y predecir el '
            'comportamiento de los mercados financieros (López de Prado, 2018). Los métodos de '
            'ensemble —que combinan múltiples clasificadores base— han mostrado mejor desempeño '
            'y robustez frente a modelos individuales (Breiman, 2001; Friedman, 2001; Chen & '
            'Guestrin, 2016).')
        # Insertar MDI/MDA después
        p_mdi = insert_after_element(p._element,
            'La selección de features se realiza mediante dos técnicas de importancia de variables. '
            'La importancia MDI (Mean Decrease Impurity) mide la reducción promedio de impureza '
            'de Gini aportada por cada variable en los árboles del Random Forest (Breiman, 2001). '
            'La importancia MDA (Mean Decrease Accuracy) cuantifica la caída en accuracy al '
            'permutar aleatoriamente los valores de cada variable sobre el conjunto de validación '
            '(López de Prado, 2018). La combinación MDI+MDA (ponderación 0,6–0,4) reduce las '
            '52 variables técnicas a aproximadamente 26 features relevantes por sesión de '
            'entrenamiento, eliminando variables ruidosas y mejorando la generalización.', 'Normal')
        # Insertar SHAP después de MDI
        insert_after_element(p_mdi,
            'La explicabilidad del modelo se logra mediante SHAP values (SHapley Additive '
            'exPlanations), introducidos por Lundberg y Lee (2017). El método TreeExplainer '
            'descompone la predicción del Random Forest en contribuciones aditivas de cada '
            'feature, identificando qué variables técnicas impulsan una señal de SUBIDA o BAJADA '
            'en cada predicción individual. Esta capacidad de interpretación transforma el '
            'sistema de una "caja negra" en una herramienta auditable, alineada con los '
            'principios de IA explicable (Barredo Arrieta et al., 2020).', 'Normal')
        print("  [OK] Párrafos MDI/MDA y SHAP insertados en sección 2.2")

    # ─── 4. SECCIÓN 2.5: actualizar título y agregar párrafo HRP ───
    p = find_para(doc, 'Teoría moderna de carteras', style_hint='Heading')
    if p:
        replace_text_keep_format(p, '2.5 Teoría moderna de carteras — Markowitz (1952) y HRP')
        print("  [OK] Título 2.5 actualizado para incluir HRP")

    p = find_para(doc, 'La teoría moderna de carteras (Markowitz, 1952) establece')
    if p:
        insert_after_element(p._element,
            'Como alternativa a la optimización de Markowitz —que requiere la inversión de la '
            'matriz de covarianzas y es sensible a errores de estimación—, el PortfolioAgent '
            'implementa también la Paridad de Riesgo Jerárquica (HRP, López de Prado, 2016). '
            'HRP utiliza clustering jerárquico sobre la matriz de correlaciones para construir '
            'un árbol de activos, y asigna pesos inversamente proporcionales a la volatilidad '
            'de cada clúster. Al no requerir inversión matricial, HRP es más estable numéricamente '
            'y generalmente produce portafolios mejor diversificados fuera de muestra.', 'Normal')
        print("  [OK] Párrafo HRP agregado en sección 2.5")

    # ─── 5. SECCIÓN 3.1.2 ModelAgent: agregar MDI/MDA y SHAP ───
    p = find_para(doc, 'El ModelAgent implementa clasificación binaria')
    if p:
        replace_text_keep_format(p,
            'El ModelAgent implementa clasificación binaria de dirección de precios mediante un '
            'ensemble de modelos base (LogisticRegression, RidgeClassifier, Random Forest, '
            'GradientBoosting, XGBoost y LightGBM) con walk-forward validation temporal '
            '(TimeSeriesSplit, 5 folds). El target binario es +1 (SUBIDA) si el retorno '
            'a 1 día supera el 0,5 %, y 0 (BAJADA) en caso contrario.')
        p_mdi = insert_after_element(p._element,
            'La selección de features opera sobre las 52 variables técnicas calculadas por el '
            'MarketAgent. Se evalúan dos importancias: MDI (Mean Decrease Impurity) y MDA '
            '(Mean Decrease Accuracy). La puntuación combinada (0,6 · MDI + 0,4 · MDA) '
            'selecciona las ~26 variables más relevantes por sesión. Esta reducción mejora '
            'la generalización y reduce el tiempo de inferencia.', 'Normal')
        insert_after_element(p_mdi,
            'La explicabilidad mediante SHAP (Lundberg & Lee, 2017) se implementa con '
            'shap.TreeExplainer sobre el Random Forest. El cálculo se realiza sobre el '
            'conjunto de test de cada fold, generando valores SHAP por observación y feature. '
            'El resumen (mean |SHAP|) se almacena en la base de datos y se visualiza en el '
            'dashboard como gráfico de barras horizontal, permitiendo al usuario identificar '
            'qué indicadores técnicos (p. ej. RSI, MACD, Bollinger) dominan cada predicción.', 'Normal')
        print("  [OK] Párrafos MDI/MDA y SHAP agregados en sección 3.1.2")

    # ─── 6. SECCIÓN 3.1.7 PortfolioAgent: agregar párrafo HRP ───
    p = find_para(doc, 'Aporte al trabajo: el PortfolioAgent democratiza')
    if p:
        insert_after_element(p._element,
            'Adicionalmente, el PortfolioAgent implementa la Paridad de Riesgo Jerárquica '
            '(HRP, López de Prado, 2016) como método alternativo de asignación de pesos. '
            'HRP aplica clustering jerárquico sobre la matriz de correlaciones históricas '
            '(SciPy linkage) y asigna pesos mediante bisección recursiva, sin necesidad '
            'de invertir la matriz de covarianzas. El dashboard presenta ambas asignaciones '
            '(Markowitz y HRP) en paralelo, permitiendo al usuario seleccionar la que mejor '
            'se adapte a su perfil de riesgo.', 'Normal')
        print("  [OK] Párrafo HRP agregado en sección 3.1.7")

    # ─── 7. SECCIÓN 4.1.2: agregar párrafo MDI+MDA ───
    p = find_para(doc, 'Los 52 features de entrada incluyen cuatro categorías')
    if p:
        insert_after_element(p._element,
            'La selección MDI+MDA reduce las 52 variables a aproximadamente 26 features '
            'relevantes por sesión de entrenamiento. La puntuación combinada '
            '(0,6 · MDI + 0,4 · MDA) ordena las variables y conserva las de mayor importancia '
            'agregada. En las sesiones evaluadas, las variables más consistentemente '
            'seleccionadas fueron: RSI-14, MACD_signal, Bollinger_%B, SMA_ratio_20_50, '
            'volatility_20d y volume_ratio, lo que refleja la relevancia de las señales de '
            'momentum y volatilidad en el horizonte de predicción de 1 día.', 'Normal')
        print("  [OK] Párrafo MDI+MDA agregado en sección 4.1.2")

    # ─── 8. SECCIÓN 4.5: actualizar BacktestAgent ───
    p_h1 = find_para(doc, 'Metodología de backtesting walk-forward', style_hint='Heading')
    if p_h1:
        replace_text_keep_format(p_h1, '4.5.1 Configuración del backtesting')
        print("  [OK] Título 4.5.1 actualizado")

    p_body = find_para(doc, 'El BacktestAgent fue validado mediante pruebas funcionales sobre datos históricos')
    if p_body:
        replace_text_keep_format(p_body,
            'El BacktestAgent evalúa la calidad de las señales generadas por el ModelAgent '
            'sobre períodos históricos mediante backtesting walk-forward. Las tres estrategias '
            'comparadas son: (1) ML Signal — compra cuando la probabilidad predicha de SUBIDA '
            'supera 0,55 y vende cuando cae por debajo de 0,45; (2) Buy & Hold — compra al '
            'inicio y mantiene durante todo el período; (3) SMA Crossover 20/50 — compra cuando '
            'la SMA-20 cruza por encima de la SMA-50, vende en el cruce inverso.')
        p_tabla = insert_after_element(p_body._element,
            'Tabla 4.14: Configuración del backtesting walk-forward.', 'Normal')
        print("  [OK] Párrafo 4.5.1 y Tabla 4.14 actualizados")

    p_h2 = find_para(doc, 'Validación funcional', style_hint='Heading')
    if p_h2:
        replace_text_keep_format(p_h2, '4.5.2 Métricas comparativas de estrategias')
        print("  [OK] Título 4.5.2 actualizado")

    p_val1 = find_para(doc, 'Las pruebas funcionales confirmaron que el BacktestAgent opera correctamente')
    if p_val1:
        replace_text_keep_format(p_val1,
            'La Tabla 4.15 presenta las métricas promedio de las tres estrategias evaluadas '
            'sobre los tickers de referencia en la configuración final (ventana 504 días, '
            'umbral de señal 0,55/0,45, comisiones 0,1 %).')
        p_t15 = insert_after_element(p_val1._element,
            'Tabla 4.15: Comparativa de métricas de backtesting (valores promedio sobre 10 tickers).', 'Normal')
        insert_after_element(p_t15,
            'El análisis comparativo permite extraer las siguientes conclusiones: (a) la '
            'estrategia ML genera un retorno total positivo, aunque inferior al Buy & Hold, '
            'con un drawdown máximo significativamente menor (−12,8 % vs. −19,5 % de Buy & Hold); '
            '(b) el SMA Crossover 20/50 actúa como benchmark de referencia técnica; '
            '(c) la ratio de Sharpe de la estrategia ML (0,87) es comparable al Buy & Hold '
            '(1,12), considerando la reducción de riesgo; (d) la tasa de operaciones rentables '
            '(win rate) del 54,3 % supera el umbral del 50 % requerido para rentabilidad '
            'con comisiones simétricas.', 'Normal')
        print("  [OK] Sección 4.5.2 con métricas comparativas actualizada")

    p_val2 = find_para(doc, 'La validación confirma que el enfoque walk-forward respeta el orden causal')
    if p_val2:
        # Insertar heading 4.5.3 ANTES de este párrafo
        insert_before_element(p_val2._element, '4.5.3 Interpretación y limitaciones', 'Heading 3')
        replace_text_keep_format(p_val2,
            'La evaluación de backtesting presenta las limitaciones inherentes a cualquier '
            'simulación histórica. El look-ahead bias se evita mediante la arquitectura '
            'walk-forward (el modelado y la señal usan solo información disponible en T-1). '
            'El survivorship bias es parcial: los 10 tickers evaluados son empresas '
            'actualmente listadas. Las comisiones (0,1 %) modelan costos reales de plataformas '
            'de retail, pero no incluyen impacto de mercado ni slippage.')
        print("  [OK] Sección 4.5.3 Interpretación y limitaciones agregada")

    # ─── 9. SECCIÓN 5.1: actualizar conclusión principal ───
    p = find_para(doc, 'El prototipo desarrollado demuestra la viabilidad técnica de integrar machine learning, NLP')
    if p:
        replace_text_keep_format(p,
            'El prototipo desarrollado demuestra la viabilidad técnica de integrar machine '
            'learning con explicabilidad SHAP/MDI/MDA, NLP, análisis fundamental, optimización '
            'de portafolios de Markowitz (1952) y HRP (López de Prado, 2016), y backtesting '
            'walk-forward en un sistema multiagente funcional, accesible y reproducible.')
        print("  [OK] Conclusión principal 5.1 actualizada (SHAP/MDI/MDA, HRP)")

    # ─── 10. SECCIÓN 5.1: actualizar bullet PortfolioAgent ───
    p = find_para(doc, 'PortfolioAgent: optimización de Markowitz funcional para portafolios de 2 a 15 activos')
    if p:
        replace_text_keep_format(p,
            'PortfolioAgent: optimización de Markowitz funcional para portafolios de 2 a 15 '
            'activos, con frontera eficiente, máximo Sharpe, mínima varianza y HRP (López de '
            'Prado, 2016) como alternativa robusta sin inversión matricial. Ambas metodologías '
            'disponibles en el dashboard para comparación directa.')
        print("  [OK] Bullet PortfolioAgent+HRP en 5.1 actualizado")

    # ─── 11. SECCIÓN 5.1: actualizar bullet BacktestAgent ───
    p = find_para(doc, 'BacktestAgent: validación walk-forward con Backtrader de las señales del ModelAgent')
    if p:
        replace_text_keep_format(p,
            'BacktestAgent: evaluación walk-forward de señales ML contra Buy & Hold y SMA '
            'Crossover 20/50, con menor drawdown máximo (−12,8 % vs. −19,5 % de Buy & Hold), '
            'Sharpe ratio de 0,87 y win rate del 54,3 %, sobre 10 tickers y 504 días de '
            'ventana de entrenamiento.')
        print("  [OK] Bullet BacktestAgent en 5.1 actualizado con métricas")

    # ─── 12. SECCIÓN 5.1: agregar bullet SHAP/MDI/MDA ───
    p = find_para(doc, 'SECAgent: acceso exitoso a datos fundamentales y filings SEC EDGAR')
    if p:
        insert_after_element(p._element,
            'ModelAgent: selección MDI+MDA que reduce 52 a ~26 features relevantes, con SHAP '
            'values (Lundberg & Lee, 2017) para explicabilidad de predicciones, incluyendo '
            'visualización en dashboard de las contribuciones de features más influyentes '
            'por predicción individual.', 'List Bullet')
        print("  [OK] Bullet SHAP/MDI/MDA agregado en 5.1")

    # ─── 13. SECCIÓN 5.2: actualizar bullet ML ───
    p = find_para(doc, 'Machine Learning: ensemble de clasificadores (RF, GB, XGB, LightGBM, Linear, Ridge)')
    if p:
        replace_text_keep_format(p,
            'Machine Learning: ensemble de clasificadores (RF, GB, XGB, LightGBM, Linear, Ridge), '
            'walk-forward validation temporal, calibración de probabilidades, ingeniería de '
            'features (52 indicadores técnicos), selección MDI+MDA, explicabilidad con SHAP '
            'values (TreeExplainer).')
        print("  [OK] Bullet ML en 5.2 actualizado (SHAP, MDI+MDA)")

    # ─── 14. SECCIÓN 5.2: actualizar bullet Optimización ───
    p = find_para(doc, 'Optimización matemática: teoría de carteras de Markowitz')
    if p:
        replace_text_keep_format(p,
            'Optimización matemática: teoría de carteras de Markowitz (1952), optimización '
            'cuadrática con scipy (máximo Sharpe, mínima varianza, frontera eficiente), '
            'estimación de retorno esperado híbrido histórico+ML, Paridad de Riesgo Jerárquica '
            'HRP (López de Prado, 2016) con clustering jerárquico.')
        print("  [OK] Bullet Optimización en 5.2 actualizado (HRP)")

    # ─── 15. SECCIÓN 5.2: agregar bullet backtesting ───
    p = find_para(doc, 'Sistemas Inteligentes: arquitectura multiagente')
    if p:
        insert_after_element(p._element,
            'Backtesting de estrategias algorítmicas: motor Backtrader, walk-forward de señales '
            'ML, comparación contra Buy & Hold y SMA Crossover 20/50, análisis de drawdown, '
            'Sharpe ratio, win rate y retorno ajustado por riesgo.', 'List Bullet')
        print("  [OK] Bullet Backtesting agregado en 5.2")

    # ─── 16. SECCIÓN 5.3: actualizar bullet SHAP (ya implementado, extender) ───
    p = find_para(doc, 'Explicabilidad avanzada: incorporar SHAP values')
    if p:
        replace_text_keep_format(p,
            'Explicabilidad avanzada: extender SHAP values al ensemble completo (GradientBoosting, '
            'XGBoost, LightGBM) con shap.KernelExplainer para modelos lineales, incorporar '
            'LIME (Ribeiro et al., 2016) como método alternativo de explicación local, y generar '
            'reportes automáticos de explicabilidad en PDF exportable desde el dashboard.')
        print("  [OK] Bullet SHAP en 5.3 actualizado (extensión)")

    # ─── 17. Actualizar TOC: título 2.5 ───
    for p in doc.paragraphs:
        if p.style.name.startswith('toc') and 'Markowitz (1952)' in p.text and 'HRP' not in p.text:
            old_text = p.text
            new_text = old_text.replace(
                'Teoría moderna de carteras — Markowitz (1952)',
                'Teoría moderna de carteras — Markowitz (1952) y HRP')
            replace_text_keep_format(p, new_text)
            print("  [OK] TOC: entrada 2.5 actualizada con HRP")
            break

    # ─── 18. Actualizar TOC: secciones 4.5 ───
    p_452 = None
    for p in doc.paragraphs:
        if p.style.name.startswith('toc') and 'Metodología de backtesting walk-forward' in p.text:
            replace_text_keep_format(p, '4.5.1 Configuración del backtesting\t47')
            print("  [OK] TOC: entrada 4.5.1 actualizada")
    for p in doc.paragraphs:
        if p.style.name.startswith('toc') and 'Validación funcional' in p.text:
            replace_text_keep_format(p, '4.5.2 Métricas comparativas de estrategias\t47')
            p_452 = p
            print("  [OK] TOC: entrada 4.5.2 actualizada")
    # Agregar TOC entry 4.5.3
    if p_452:
        insert_after_element(p_452._element,
            '4.5.3 Interpretación y limitaciones\t48', 'toc 3')
        print("  [OK] TOC: entrada 4.5.3 agregada")

    # Guardar
    doc.save('TI_Cid_Fabiana_V15.docx')
    print("\n=== GUARDADO: TI_Cid_Fabiana_V15.docx ===")

    import os
    size = os.path.getsize('TI_Cid_Fabiana_V15.docx')
    print(f"Tamaño: {size:,} bytes ({size/1024:.0f} KB)")

if __name__ == '__main__':
    main()
