"""
Genera el informe de estado del proyecto para el director de tesis.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color='1F3864'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_paragraph(doc, text, bold=False, italic=False, size=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_table_row(table, col1, col2, header=False, bg=None):
    row = table.add_row()
    row.cells[0].text = col1
    row.cells[1].text = col2
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            for run in para.runs:
                run.font.size = Pt(9)
                if header:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if bg:
            set_cell_bg(cell, bg)
    return row

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── ENCABEZADO ───
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('INFORME DE AVANCE — TESIS FINAL DE MAESTRÍA')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string('1F3864')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run('Sistema Multiagente de Análisis y Optimización de Portafolios Financieros')
run2.bold = True
run2.font.size = Pt(11)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(10)
run3 = p3.add_run('Ing. María Fabiana Cid  |  Maestría en Finanzas — UNR  |  Junio 2026')
run3.font.size = Pt(9)
run3.italic = True

# Línea separadora
doc.add_paragraph('─' * 90).paragraph_format.space_after = Pt(6)

# ─── 1. IDENTIFICACIÓN ───
add_heading(doc, '1. Identificación del trabajo', level=2)

tabla_id = doc.add_table(rows=0, cols=2)
tabla_id.style = 'Table Grid'
tabla_id.columns[0].width = Cm(4.5)
tabla_id.columns[1].width = Cm(11.5)

datos = [
    ('Tesista',        'Ing. María Fabiana Cid'),
    ('Director',       'Ph.D. Luciano Machain (UNR)'),
    ('Programa',       'Maestría en Finanzas — Universidad Nacional de Rosario'),
    ('Versión actual', 'V15 — junio 2026  (53+ páginas)'),
    ('Repositorio',    'github.com/fabianacid/TP_Final  (rama: maestria_tesis)'),
    ('Stack principal','Python 3.10 · FastAPI · Streamlit · scikit-learn · yfinance · SEC EDGAR'),
]
for k, v in datos:
    row = tabla_id.add_row()
    row.cells[0].text = k
    row.cells[1].text = v
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            for run in para.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
        set_cell_bg(cell, 'EBF0FA' if i == 0 else 'FFFFFF')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── 2. ESTADO DE AVANCE ───
add_heading(doc, '2. Estado de avance', level=2)
add_paragraph(doc,
    'El prototipo funcional se encuentra completamente implementado y validado empíricamente. '
    'Los ocho agentes especializados están operativos, la API REST y el dashboard Streamlit '
    'están integrados, y se han ejecutado cuatro sesiones de prueba con datos reales de mercado.',
    size=10, space_after=4)

tabla_av = doc.add_table(rows=0, cols=3)
tabla_av.style = 'Table Grid'
tabla_av.columns[0].width = Cm(6.5)
tabla_av.columns[1].width = Cm(5.0)
tabla_av.columns[2].width = Cm(4.5)

header_row = tabla_av.add_row()
for i, txt in enumerate(['Componente', 'Estado', 'Validación']):
    header_row.cells[i].text = txt
    for para in header_row.cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(header_row.cells[i], '1F3864')

filas = [
    ('MarketAgent — 52 indicadores técnicos',      '✅ Implementado', 'Datos reales Yahoo Finance'),
    ('ModelAgent — Ensemble ML + SHAP/MDI/MDA',   '✅ Implementado', 'Walk-forward 5 folds, 10 tickers'),
    ('SentimentAgent — Ensemble NLP (FinBERT+3)',  '✅ Implementado', 'Noticias reales + filings 8-K'),
    ('RecommendationAgent — multi-factor 15+ vars','✅ Implementado', '30 pruebas funcionales'),
    ('AlertAgent — 5 técnicas detección anomalías','✅ Implementado', '6 niveles de severidad'),
    ('SECAgent — yfinance + SEC EDGAR API',        '✅ Implementado', 'Ratios + filings 10 tickers'),
    ('PortfolioAgent — Markowitz + HRP',           '✅ Implementado', 'Frontera eficiente, Sharpe, VaR'),
    ('BacktestAgent — walk-forward Backtrader',    '✅ Implementado', 'ML vs Buy&Hold vs SMA 20/50'),
    ('API REST FastAPI + autenticación JWT/bcrypt', '✅ Implementado', 'OWASP / Comunicación A 7724'),
    ('Dashboard Streamlit — 4 pestañas',           '✅ Implementado', 'Capturas en Capítulo 4'),
]
for comp, estado, val in filas:
    row = tabla_av.add_row()
    row.cells[0].text = comp
    row.cells[1].text = estado
    row.cells[2].text = val
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.space_before = Pt(1)
            for run in para.runs:
                run.font.size = Pt(9)
        bg = 'F2F8F2' if i == 1 else ('FFFFFF' if i != 0 else 'F7F9FC')
        set_cell_bg(cell, bg)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── 3. RESULTADOS EMPÍRICOS ───
add_heading(doc, '3. Resultados empíricos clave', level=2)

tabla_res = doc.add_table(rows=0, cols=4)
tabla_res.style = 'Table Grid'
for i, w in enumerate([Cm(5.5), Cm(3.5), Cm(3.5), Cm(3.5)]):
    tabla_res.columns[i].width = w

header2 = tabla_res.add_row()
for i, txt in enumerate(['Métrica', 'Resultado', 'Requisito', 'Estado']):
    header2.cells[i].text = txt
    for para in header2.cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(header2.cells[i], '1F3864')

metricas = [
    ('Tasa de éxito funcional',     '100 % (30/30 pruebas)', '≥ 95 %',     '✅ Cumple'),
    ('Latencia promedio',           '4,65 s',                '< 5 s',      '✅ Cumple'),
    ('Accuracy ensemble ML',        '57,0 % (rango 53–63 %)', '> 50 %',   '✅ Cumple'),
    ('AUC-ROC',                     '0,579 promedio',        '> 0,5',      '✅ Cumple'),
    ('Usuarios concurrentes',       '10 (100 % éxito)',      'N/D',        '⚠ Single-worker'),
    ('Drawdown ML vs Buy & Hold',   '−12,8 % vs −19,5 %',   'Menor riesgo','✅ Cumple'),
    ('Sharpe ratio ML',             '0,87',                  '> 0',        '✅ Cumple'),
    ('Win rate backtesting',        '54,3 %',                '> 50 %',     '✅ Cumple'),
    ('Disponibilidad (1–10 users)', '100 %',                 '≥ 99 %',     '✅ Cumple'),
]
for m, r, req, est in metricas:
    row = tabla_res.add_row()
    row.cells[0].text = m
    row.cells[1].text = r
    row.cells[2].text = req
    row.cells[3].text = est
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.space_before = Pt(1)
            for run in para.runs:
                run.font.size = Pt(9)
        bg = 'F2F8F2' if '✅' in row.cells[3].text else ('FFF8F0' if '⚠' in row.cells[3].text else 'FFFFFF')
        set_cell_bg(cell, bg if i == 3 else 'FFFFFF')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── 4. DETALLE DE LOS AGENTES ───
add_heading(doc, '4. Arquitectura multiagente — detalle de los ocho agentes', level=2)
add_paragraph(doc,
    'El sistema se organiza en ocho agentes especializados que cooperan mediante un pipeline '
    'secuencial y paralelo coordinado por la API FastAPI. Cada agente es autónomo, con '
    'responsabilidades bien delimitadas y salidas estructuradas que alimentan al siguiente.',
    size=10, space_after=4)

tabla_ag = doc.add_table(rows=0, cols=3)
tabla_ag.style = 'Table Grid'
tabla_ag.columns[0].width = Cm(3.2)
tabla_ag.columns[1].width = Cm(5.8)
tabla_ag.columns[2].width = Cm(7.0)

hdr = tabla_ag.add_row()
for i, txt in enumerate(['Agente', 'Técnicas principales', 'Salida al sistema']):
    hdr.cells[i].text = txt
    for para in hdr.cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(hdr.cells[i], '1F3864')

agentes = [
    ('MarketAgent',
     '52 indicadores técnicos: tendencia (SMA, EMA, MACD, ADX), momentum (RSI, Stochastic), '
     'volatilidad (Bollinger Bands, ATR) y volumen (OBV, VWAP). Datos via yfinance.',
     'Serie histórica enriquecida con features técnicas. Fuente de datos para todos los agentes.'),

    ('ModelAgent',
     'Ensemble de 6 clasificadores (Random Forest, GradientBoosting, XGBoost, LightGBM, '
     'LogisticRegression, RidgeClassifier). Walk-forward validation temporal (5 folds, '
     'TimeSeriesSplit). Selección de features MDI+MDA (52 → ~26 variables). '
     'SHAP values (TreeExplainer) para explicabilidad.',
     'Predicción binaria SUBIDA/BAJADA con probabilidad calibrada. '
     'Accuracy 57 %, AUC-ROC 0,579. Gráfico SHAP de importancia de features en dashboard.'),

    ('SentimentAgent',
     'Ensemble NLP ponderado: FinBERT (40 %), VADER (25 %), lexicón financiero 500+ términos (20 %), '
     'TextBlob (15 %). Fuentes: noticias Yahoo Finance y filings 8-K de SEC EDGAR.',
     'Score de sentimiento [−1, +1] y categoría (Positivo/Neutral/Negativo) '
     'por ticker. Tendencia temporal del sentimiento.'),

    ('RecommendationAgent',
     'Sistema de decisión multi-factor con 15+ variables ponderadas por categoría: '
     'señal ML (40 %), sentimiento NLP (25 %), score fundamental SEC (20 %), '
     'indicadores técnicos (15 %).',
     'Recomendación textual accionable (Comprar / Mantener / Reducir / Vender) '
     'con justificación explicativa apta para inversores sin formación financiera.'),

    ('AlertAgent',
     '5 técnicas de detección de anomalías: Z-Score, MAD (Median Absolute Deviation), '
     'CUSUM (cambios de tendencia), EWMA (control estadístico) e Isolation Forest '
     '(anomalías multivariadas). Rate limiting inteligente.',
     '6 niveles de severidad: INFO, LOW, MEDIUM, HIGH, CRITICAL, EMERGENCY. '
     'Alertas contextuales con explicación y acción sugerida.'),

    ('SECAgent',
     'Dos fuentes oficiales gratuitas: yfinance.info para ratios en tiempo real '
     '(P/E, P/B, ROE, ROA, márgenes, deuda/equity) y API SEC EDGAR para filings '
     'recientes (8-K, 10-Q, 10-K). Caché de 2 horas.',
     'Score fundamental multi-factor [0–100] y resumen de ratios financieros '
     'clave. Acceso sin costo de licencia para todo el mercado estadounidense.'),

    ('PortfolioAgent',
     'Ejecución paralela del pipeline completo por activo (ThreadPoolExecutor). '
     'Optimización de Markowitz (máximo Sharpe, mínima varianza, frontera eficiente) '
     'vía scipy. HRP (López de Prado, 2016): clustering jerárquico + bisección recursiva '
     'sin inversión matricial. Retorno esperado híbrido: 70 % histórico + 30 % ML.',
     'Métricas del portafolio actual (retorno, volatilidad, Sharpe, VaR 95 %). '
     'Pesos óptimos Markowitz y HRP en paralelo. Alertas de concentración y correlación.'),

    ('BacktestAgent',
     'Motor Backtrader con walk-forward sobre datos históricos. 3 estrategias: '
     'ML Signal (umbral probabilidad 0,55/0,45), Buy & Hold y SMA Crossover 20/50. '
     'Comisiones 0,1 %. Capital inicial USD 100.000.',
     'Comparativa de retorno total, drawdown máximo, Sharpe ratio y win rate. '
     'Resultado: ML −12,8 % drawdown vs −19,5 % B&H. Sharpe 0,87. Win rate 54,3 %.'),
]

for nombre, tecnicas, salida in agentes:
    row = tabla_ag.add_row()
    # Col 0: nombre en negrita
    p0 = row.cells[0].paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    p0.paragraph_format.space_before = Pt(2)
    r0 = p0.add_run(nombre)
    r0.bold = True
    r0.font.size = Pt(9)
    set_cell_bg(row.cells[0], 'EBF0FA')
    # Col 1 y 2: texto normal
    for ci, txt in [(1, tecnicas), (2, salida)]:
        p = row.cells[ci].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(txt)
        r.font.size = Pt(8.5)
        set_cell_bg(row.cells[ci], 'FFFFFF')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── 5. INCORPORACIONES RECIENTES (V15) ───
add_heading(doc, '5. Incorporaciones recientes — versión 15', level=2)

items_v15 = [
    ('Explicabilidad SHAP/MDI/MDA',
     'El ModelAgent incorpora selección de features MDI+MDA (reduce 52 a ~26 variables relevantes) '
     'y SHAP values con TreeExplainer para interpretar cada predicción individual. '
     'Visualización integrada en el dashboard como gráfico de importancia de features.'),
    ('Paridad de Riesgo Jerárquica — HRP',
     'El PortfolioAgent implementa HRP (López de Prado, 2016) como alternativa a Markowitz. '
     'Usa clustering jerárquico sobre la matriz de correlaciones, eliminando la necesidad de '
     'inversión matricial y mejorando la estabilidad numérica fuera de muestra. '
     'El dashboard presenta Markowitz y HRP en paralelo para comparación directa.'),
    ('BacktestAgent — evaluación comparativa mejorada',
     'La sección 4.5 fue reorganizada en tres subsecciones: Configuración del backtesting (4.5.1), '
     'Métricas comparativas de las tres estrategias con Tabla 4.15 (4.5.2), e Interpretación '
     'y limitaciones del backtesting histórico (4.5.3).'),
    ('Marco teórico actualizado',
     'Secciones 2.2 y 2.5 ampliadas con fundamentos de MDI/MDA, SHAP y HRP. '
     'Conclusiones (5.1) y aplicación de contenidos del posgrado (5.2) actualizadas.'),
]

for titulo, desc in items_v15:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    run_b = p.add_run(titulo + ': ')
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_n = p.add_run(desc)
    run_n.font.size = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── 6. PRÓXIMOS PASOS ───
add_heading(doc, '6. Trabajo pendiente y próximos pasos', level=2)

pendientes = [
    'Corrección y ajuste fino de la memoria técnica V15 según observaciones del director.',
    'Extensión de SHAP values al ensemble completo (XGBoost, LightGBM) para análisis comparativo de explicabilidad entre modelos.',
    'Incorporación de al menos un feature macroeconómico (VIX o tasa de referencia) como variable exógena en el ModelAgent.',
    'Preparación de la presentación final y defensa oral ante el jurado.',
]
for item in pendientes:
    add_bullet(doc, item, size=10)

# ─── PIE DE PÁGINA ───
doc.add_paragraph('─' * 90).paragraph_format.space_before = Pt(8)
p_pie = doc.add_paragraph()
p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_pie.paragraph_format.space_after = Pt(0)
run_pie = p_pie.add_run(
    'Documento generado el 19 de junio de 2026  ·  '
    'Repositorio: github.com/fabianacid/TP_Final  ·  '
    'Versión memoria: TI_Cid_Fabiana_V15.docx')
run_pie.font.size = Pt(8)
run_pie.italic = True
run_pie.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save('Informe_Avance_Director_Junio2026_v2.docx')
print("Guardado: Informe_Avance_Director_Junio2026_v2.docx")

import os
size = os.path.getsize('Informe_Avance_Director_Junio2026_v2.docx')
print(f"Tamaño: {size:,} bytes ({size/1024:.0f} KB)")
